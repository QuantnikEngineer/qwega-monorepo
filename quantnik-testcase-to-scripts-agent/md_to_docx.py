"""
Convert Langfuse_Instrumentation_Documentation.md to .docx format.
One-time utility — delete after use.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

MD_PATH = Path(__file__).resolve().parent / "Langfuse_Instrumentation_Documentation.md"
DOCX_PATH = Path(__file__).resolve().parent / "Langfuse_Instrumentation_Documentation.docx"


def _add_formatted_run(paragraph, text):
    """Parse inline markdown (bold, code, links) and add runs to paragraph."""
    # Pattern for bold, inline code, and markdown links
    pattern = r'(\*\*(.+?)\*\*)|(`(.+?)`)|(\[(.+?)\]\((.+?)\))'
    last_end = 0
    for m in re.finditer(pattern, text):
        # Add plain text before this match
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])
        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(4):  # inline code
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif m.group(6):  # link - just show the text
            run = paragraph.add_run(m.group(6))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        last_end = m.end()
    # Remaining plain text
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def convert():
    md_text = MD_PATH.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Heading styles
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            # Add a thin line via a paragraph with bottom border
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Headings
        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            level = min(level, 4)
            heading_text = stripped.lstrip('#').strip()
            h = doc.add_heading(level=level)
            _add_formatted_run(h, heading_text)
            i += 1
            continue

        # Code blocks (fenced)
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Tables
        if stripped.startswith('|') and '|' in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            # Parse table
            if len(table_lines) < 2:
                continue

            # Remove separator row (second line with dashes)
            header_row = [c.strip() for c in table_lines[0].split('|')[1:-1]]
            data_rows = []
            for tl in table_lines[2:]:  # skip header and separator
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                data_rows.append(cells)

            num_cols = len(header_row)
            table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Header
            for j, h in enumerate(header_row):
                cell = table.rows[0].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                _add_formatted_run(p, h)
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)

            # Data rows
            for r_idx, row_data in enumerate(data_rows):
                for c_idx in range(min(len(row_data), num_cols)):
                    cell = table.rows[r_idx + 1].cells[c_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    _add_formatted_run(p, row_data[c_idx])
                    for run in p.runs:
                        run.font.size = Pt(10)

            doc.add_paragraph()  # spacing after table
            continue

        # Bullet points
        if stripped.startswith('- '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_run(p, text)
            i += 1
            continue

        # Numbered list items (e.g., "1. ", "2. ")
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_run(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_run(p, stripped)
        i += 1

    doc.save(str(DOCX_PATH))
    print(f"✅ DOCX saved: {DOCX_PATH}")


if __name__ == "__main__":
    convert()
