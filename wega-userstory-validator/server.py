import logging
import re
from io import BytesIO
from flask import Flask, request, jsonify
from flask_cors import CORS
from flasgger import Swagger, swag_from
from google import genai
from google.genai import types
import os
from dotenv import load_dotenv

load_dotenv()
import ssl
import base64
import requests as http_requests
from docx import Document
from xhtml2pdf import pisa
from html.parser import HTMLParser

ssl._create_default_https_context = ssl._create_unverified_context

import httpx
_original_httpx_client_init = httpx.Client.__init__
def _patched_httpx_client_init(self, *args, **kwargs):
    kwargs.setdefault("verify", False)
    _original_httpx_client_init(self, *args, **kwargs)
httpx.Client.__init__ = _patched_httpx_client_init

_original_httpx_async_client_init = httpx.AsyncClient.__init__
def _patched_httpx_async_client_init(self, *args, **kwargs):
    kwargs.setdefault("verify", False)
    _original_httpx_async_client_init(self, *args, **kwargs)
httpx.AsyncClient.__init__ = _patched_httpx_async_client_init

app = Flask(__name__)

# Enable CORS for all routes
CORS(app, resources={r"/*": {"origins": "*"}})

# Swagger configuration
swagger_config = {
    "headers": [],
    "specs": [
        {
            "endpoint": "apispec",
            "route": "/apispec.json",
            "rule_filter": lambda rule: True,
            "model_filter": lambda tag: True,
        }
    ],
    "static_url_path": "/flasgger_static",
    "swagger_ui": True,
    "specs_route": "/docs"
}

swagger_template = {
    "info": {
        "title": "Wega Userstory Validator API",
        "description": "API for analyzing Business Requirements Documents (BRD) against User Stories using Vertex AI (Gemini)",
        "version": "1.0.0",
        "contact": {
            "name": "Wega Validator Team"
        }
    },
    "basePath": "/",
    "schemes": ["https", "http"]
}

swagger = Swagger(app, config=swagger_config, template=swagger_template)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

logger.info("Starting Wega Userstory Validator server")

def extract_page_id(confluence_link):
  """Extract page ID from a Confluence URL."""
  match = re.search(r'/pages/(\d+)', confluence_link)
  if match:
    return match.group(1)
  raise ValueError(f"Could not extract page ID from Confluence link: {confluence_link}")

def extract_base_url(confluence_link):
  """Extract base URL (scheme + host) from a Confluence URL."""
  match = re.match(r'(https?://[^/]+)', confluence_link)
  if match:
    return match.group(1)
  raise ValueError(f"Could not extract base URL from Confluence link: {confluence_link}")

class HTMLTextExtractor(HTMLParser):
  """Simple HTML parser to extract text content."""
  def __init__(self):
    super().__init__()
    self.text_parts = []
    
  def handle_data(self, data):
    stripped = data.strip()
    if stripped:
      self.text_parts.append(stripped)
      
  def get_text(self):
    return '\n'.join(self.text_parts)

def extract_html_text(html_content):
  """Extract plain text from HTML content."""
  parser = HTMLTextExtractor()
  parser.feed(html_content)
  return parser.get_text()

def fetch_confluence_docx(confluence_link):
  """Fetch the Confluence page content and create a PDF document.
  
  Reads the content of the specified Confluence page, converts it to PDF,
  and saves it as 'final_brd.pdf'.
  
  Args:
    confluence_link: URL of the Confluence page
    
  Returns:
    dict: Dictionary containing:
      - 'pdf_path': Path to the generated PDF file (final_brd.pdf)
      - 'text_content': Extracted plain text from the page
  """

  confluence_email = os.environ.get("CONFLUENCE_EMAIL")
  confluence_api_token = os.environ.get("CONFLUENCE_API_TOKEN")
  if not confluence_email or not confluence_api_token:
    raise ValueError("CONFLUENCE_EMAIL and CONFLUENCE_API_TOKEN environment variables are required")

  # Extract page ID and base URL from the Confluence link
  page_id = extract_page_id(confluence_link)
  base_url = extract_base_url(confluence_link)
  
  auth = (confluence_email, confluence_api_token)
  
  # Fetch the page content using Confluence REST API
  content_url = f"{base_url}/wiki/rest/api/content/{page_id}?expand=body.storage"
  logger.info(f"Fetching page content from: {content_url}")
  
  resp = http_requests.get(content_url, auth=auth, verify=False)
  resp.raise_for_status()
  
  page_data = resp.json()
  page_title = page_data.get('title', 'Confluence Page')
  page_body_html = page_data.get('body', {}).get('storage', {}).get('value', '')
  
  if not page_body_html:
    raise ValueError("No content found on the Confluence page")
  
  logger.info(f"Retrieved page content for: {page_title}")
  
  # Extract plain text from HTML for analysis
  text_content = extract_html_text(page_body_html)
  logger.info(f"Extracted text length: {len(text_content)} characters")
  
  # Create a complete HTML document with proper structure
  html_content = f"""
  <!DOCTYPE html>
  <html>
  <head>
    <meta charset="UTF-8">
    <title>{page_title}</title>
    <style>
      body {{
        font-family: Arial, sans-serif;
        margin: 20px;
        line-height: 1.6;
      }}
      h1, h2, h3, h4, h5, h6 {{
        color: #333;
        margin-top: 20px;
      }}
      table {{
        border-collapse: collapse;
        width: 100%;
        margin: 10px 0;
      }}
      th, td {{
        border: 1px solid #ddd;
        padding: 8px;
        text-align: left;
      }}
      th {{
        background-color: #f2f2f2;
      }}
    </style>
  </head>
  <body>
    <h1>{page_title}</h1>
    {page_body_html}
  </body>
  </html>
  """
  
  # Generate PDF file
  pdf_filename = "final_brd.pdf"
  logger.info(f"Generating PDF: {pdf_filename}")
  
  try:
    with open(pdf_filename, "wb") as pdf_file:
      # Convert HTML to PDF
      pisa_status = pisa.CreatePDF(
        html_content,
        dest=pdf_file
      )
    
    if pisa_status.err:
      raise RuntimeError(f"Error generating PDF: {pisa_status.err}")
    
    logger.info(f"Successfully created PDF: {pdf_filename}")
    
    return {
      'pdf_path': pdf_filename,
      'text_content': text_content
    }
    
  except Exception as e:
    logger.error(f"Error creating PDF: {str(e)}")
    raise

def extract_docx_text(docx_bytes):
  """Extract all text content (paragraphs and tables) from a .docx file."""
  doc = Document(BytesIO(docx_bytes))
  parts = []
  for para in doc.paragraphs:
    if para.text.strip():
      parts.append(para.text)
  for table in doc.tables:
    for row in table.rows:
      row_text = ' | '.join(cell.text.strip() for cell in row.cells)
      if row_text.strip():
        parts.append(row_text)
  return '\n'.join(parts)

@app.route('/health', methods=['GET'])
def health_check():
  """Health check endpoint
  ---
  tags:
    - Health
  responses:
    200:
      description: Service is healthy
      schema:
        type: object
        properties:
          status:
            type: string
            example: healthy
          service:
            type: string
            example: Wega Userstory Validator
  """
  logger.info("Health check endpoint called")
  return jsonify({"status": "healthy", "service": "Wega Userstory Validator"}), 200

@app.route('/v1/api/validate-user-story', methods=['POST'])
def analyze_brd():
  """Analyze BRD documents against User Stories
  ---
  tags:
    - Analysis
  parameters:
    - in: body
      name: body
      required: true
      schema:
        type: object
        required:
          - validate_user_story_text
          - confluence_link
        properties:
          validate_user_story_text:
            type: string
            description: The user story or feature description to analyze
            example: "User authentication with OAuth 2.0"
          confluence_link:
            type: string
            description: Confluence page URL containing the BRD content (page content will be converted to PDF)
            example: "https://wegabuildiq.atlassian.net/wiki/spaces/WAAD/pages/36962305/BRD+Page"
  responses:
    200:
      description: Analysis completed successfully
      schema:
        type: object
        properties:
          status:
            type: string
            example: success
          message:
            type: string
            example: Analysis completed
          result:
            type: string
            description: Detailed analysis with covered requirements, gaps, and recommendations
    400:
      description: Missing required fields
      schema:
        type: object
        properties:
          status:
            type: string
            example: error
          message:
            type: string
            example: validate_user_story_text and confluence_link are required
    500:
      description: Internal server error
      schema:
        type: object
        properties:
          status:
            type: string
            example: error
          message:
            type: string
  """
  logger.info("Analyze endpoint called")
  try:
    data = request.get_json()
    logger.info(f"Received request data: {data}")
    
    # Extract fields from request
    validate_user_story_text = data.get('validate_user_story_text')
    confluence_link = data.get('confluence_link')
    
    # Validate required fields
    if not validate_user_story_text or not confluence_link:
      return jsonify({
        "status": "error",
        "message": "validate_user_story_text and confluence_link are required"
      }), 400
    
    logger.info(f"validate_user_story_text: {validate_user_story_text}")
    logger.info(f"Confluence Link: {confluence_link}")
    
    # Fetch BRD document from Confluence
    logger.info("Fetching BRD document from Confluence...")
    confluence_data = fetch_confluence_docx(confluence_link)
    pdf_path = confluence_data['pdf_path']
    brd_text = confluence_data['text_content']
    logger.info(f"Generated PDF: {pdf_path}")
    logger.info(f"Extracted BRD text length: {len(brd_text)} characters")
    logger.debug(brd_text)
    # Call generate function and capture output
    generation_result = generate(validate_user_story_text, brd_text)
    jira_dictionary = generation_result.get("jira_dictionary", {})
    new_jira_createlist = generation_result.get("new_jira_createlist", [])
    result_text = generation_result.get("result_text", "")
    logger.info("Analysis completed successfully, returning from Wega-userstory-validator endpoint")

    # Extract updated_user_stories from result_text (markdown table)
    
    logger.info(f"Returning analysis result with updated_user_stories_list: {jira_dictionary}")
    return jsonify({
      "status": "success",
      "message": "Analysis completed",
      "result_text": result_text,
      "jira_dictionary": jira_dictionary,
      "new_jira_createlist": new_jira_createlist
    }), 200
    
  except ValueError as e:
    if "CONFLUENCE_EMAIL" in str(e) or "CONFLUENCE_API_TOKEN" in str(e):
      logger.error(f"server.py analyze_brd method onfluence email or token missing or has expired in user story validator agent,Confluence auth error: {str(e)}", exc_info=True)
      return jsonify({
        "status": "error",
        "message": "confluence email or token missing or has expired in user story validator agent, please check."
      }), 401
    logger.error(f"ValueError in server.py analyze_brd method analyze endpoint: {str(e)}", exc_info=True)
    return jsonify({
      "status": "error",
      "message": str(e)
    }), 400
  except Exception as e:
    logger.error(f"Error in server.py analyze_brd method endpoint: {str(e)}", exc_info=True)
    return jsonify({
      "status": "error",
      "message": str(e)
    }), 500

@app.route('/', methods=['GET'])
def home():
  """Home endpoint with service information
  ---
  tags:
    - Info
  responses:
    200:
      description: Service information
      schema:
        type: object
        properties:
          service:
            type: string
            example: Wega Userstory Validator
          version:
            type: string
            example: "1.0"
          endpoints:
            type: object
  """
  logger.info("Home endpoint called")
  return jsonify({
    "service": "Wega Userstory Validator",
    "version": "1.0",
    "endpoints": {
      "/": "Home - this page",
      "/health": "Health check",
      "/v1/api/validate-user-story": "POST - Analyze User Stories against BRD document",
      "/docs": "API Documentation (Swagger UI)"
    }
  }), 200

def generate(validate_user_story_text, brd_text):
  """Generate content using Gemini model and return structured JSON response"""
  logger.info("Starting generate() function")
  result_text = ""
  try:
    api_key = os.environ.get("GOOGLE_CLOUD_API_KEY")
    if not api_key:
      logger.warning("GOOGLE_CLOUD_API_KEY environment variable not set")
    else:
      logger.info("API key found, initializing Vertex AI client")

    client = genai.Client(
        vertexai=True,
        api_key=api_key,
        location="global",
    )
    logger.info("Vertex AI client initialized successfully")

    system_instruction = """You are an expert Business Analyst who validates and creates User Stories based on a Business Requirements Document (BRD).

You MUST respond with ONLY a single valid JSON object. No markdown, no code fences, no explanations outside the JSON.

The JSON response MUST strictly follow this schema:
{
  "summary": "<string>",
  "updated_stories_map": {
    "<JIRA_ID>": {
      "title": "<concise descriptive title>",
      "story": "<updated user story text>"
    },
    ...
  },
  "new_jira_createlist": {
    "<epic_id>": [
      {
        "title": "<concise descriptive title>",
        "story": "<new user story text>"
      },
      ...
    ],
    ...
  }
}

FIELD DEFINITIONS:
- "summary": A single string containing the high-level analysis summary. Include: which user stories are correct, partially correct, missing requirements, incorrect aspects, and BRD evidence for each point. Use \\n for line breaks within the string.
- "updated_stories_map": A JSON object where each key is the EXACT original Jira ID (e.g. "PROJ-123") from the input user stories, and each value is an object containing "title" (a concise, descriptive title for the user story) and "story" (the refined/updated user story text in proper user story format). Every Jira ID from the input MUST appear as a key in this object.
- "new_jira_createlist": A JSON object where each key is the Epic ID (epic_id) extracted from the Epic Key associated with the input user stories, and each value is an array of objects. Each object contains "title" (a concise, descriptive title) and "story" (a new user story in proper user story format) that should be created under that epic because the BRD covers requirements not addressed by any existing user story. If no new stories are needed, return an empty object {}. Every new user story MUST be mapped to an epic_id from the input.

STRICT RULES:
- Every Jira ID present in the input user stories MUST appear as a key in "updated_stories_map". Do NOT omit, rename, or alter any Jira ID.
- Do NOT use markdown formatting (no **, no ##, no ```).
- In user story text and acceptance criteria, use only \\n for newlines and - for bullet points. No other special characters.
- User stories must follow the format: "As a <role>, I want <goal>, so that <benefit>\\nAcceptance Criteria:\\n- <criterion 1>\\n- <criterion 2>\\n..."
- Every new user story in "new_jira_createlist" MUST be grouped under a valid epic_id from the input. Do NOT suggest any new user story without associating it to an epic_id.
- Return ONLY the JSON object. No text before or after it."""

    analysis_prompt = f"""Analyze the following user stories against the BRD document and produce the JSON response.

USER STORIES TO ANALYZE:
{validate_user_story_text}

MASTER BRD DOCUMENT CONTENT:
{brd_text}

Perform the following steps for EACH user story in the input:

Step 1: Identify all user stories and their Jira IDs from the USER STORIES TO ANALYZE section.

Step 2: For each user story, search the BRD content for sections that relate to that user story and identify specific requirements.

Step 3: For each user story, determine:
   a) What aspects are correct and aligned with the BRD
   b) What aspects are partially correct and need refinement
   c) What requirements from the BRD are missing in the user story
   d) What aspects are incorrect or contradict the BRD
   e) Generate a concise, descriptive title for each user story
   Cite specific BRD sections as evidence for each finding.

Step 4: Compile a high-level summary of the analysis across all user stories. Include the findings from Step 3 with BRD evidence. This becomes the "summary" field.

Step 5: For each user story, produce an updated version that:
   - Retains the EXACT original Jira ID as the key
   - Includes the generated title from Step 3e
   - Incorporates all relevant BRD requirements
   - Follows the standard user story format with acceptance criteria
   - Fixes any incorrect or missing aspects found in Step 3
   Add each as a key-value pair in "updated_stories_map" where key = exact Jira ID, value = object with "title" and "story" fields.

Step 6: Review the entire BRD for any requirements NOT covered by any of the input user stories. For each uncovered requirement, identify the most relevant Epic Key (epic_id) from the input user stories under which the new story logically belongs. Create a new user story with a concise title and proper story format, and add it as an object with "title" and "story" fields to "new_jira_createlist" grouped under the corresponding epic_id key. Do NOT create any new user story without associating it to an epic_id from the input. If no matching epic_id exists for an uncovered requirement, do NOT create a new story for it.

Step 7: Verify that:
   - ALL original Jira IDs appear in "updated_stories_map"
   - No Jira ID has been modified or omitted
   - All BRD requirements are addressed either in updated stories or new stories
   - The response is valid JSON with no markdown

Return the final JSON response now."""

    logger.info(f"Analysis prompt constructed for user story: {validate_user_story_text[:100]}...")

    content_parts = [types.Part.from_text(text=analysis_prompt)]
    model = "gemini-3-flash-preview"
    logger.info(f"Using model: {model}")

    contents = [
      types.Content(
        role="user",
        parts=content_parts
      )
    ]
    logger.info(f"Content prepared with {len(content_parts)} parts")

    generate_content_config = types.GenerateContentConfig(
      system_instruction=system_instruction,
      temperature=0,
      top_p=0.90,
      max_output_tokens=65535,
      response_mime_type="application/json",
      safety_settings=[
        types.SafetySetting(
          category="HARM_CATEGORY_HATE_SPEECH",
          threshold="BLOCK_MEDIUM_AND_ABOVE"
        ),
        types.SafetySetting(
          category="HARM_CATEGORY_DANGEROUS_CONTENT",
          threshold="BLOCK_MEDIUM_AND_ABOVE"
        ),
        types.SafetySetting(
          category="HARM_CATEGORY_SEXUALLY_EXPLICIT",
          threshold="BLOCK_MEDIUM_AND_ABOVE"
        ),
        types.SafetySetting(
          category="HARM_CATEGORY_HARASSMENT",
          threshold="BLOCK_MEDIUM_AND_ABOVE"
        ),
      ],
    )
    logger.info("Generation config: temperature=0, top_p=0.90, max_tokens=65535, response_mime_type=application/json")

    logger.info("Starting content generation stream")
    chunk_count = 0
    for chunk in client.models.generate_content_stream(
      model=model,
      contents=contents,
      config=generate_content_config,
    ):
      chunk_count += 1
      if chunk.text:
        result_text += chunk.text
      if chunk_count % 10 == 0:
        logger.debug(f"Processed {chunk_count} chunks")

    logger.debug(f"Content generation completed. Total chunks: {chunk_count}")

    import json

    result_text = result_text.strip()
    if result_text.startswith("```"):
      result_text = re.sub(r'^```(?:json)?\s*', '', result_text)
      result_text = re.sub(r'\s*```$', '', result_text)
      result_text = result_text.strip()

    try:
      parsed_response = json.loads(result_text)
    except json.JSONDecodeError as e:
      logger.error(f"Failed to parse JSON response: {e}")
      logger.debug(f"Raw response: {result_text[:1000]}")
      json_match = re.search(r'\{[\s\S]*\}', result_text)
      if json_match:
        parsed_response = json.loads(json_match.group(0))
        logger.debug("Extracted JSON from raw response via regex fallback")
      else:
        raise ValueError(f"Model did not return valid JSON: {e}")

    summary = parsed_response.get("summary", "")
    jira_dictionary = parsed_response.get("updated_stories_map", {})
    new_jira_createlist = parsed_response.get("new_jira_createlist", [])

    # Log Jira IDs with their titles for better visibility
    stories_with_titles = {k: v.get("title", "No title") if isinstance(v, dict) else "Legacy format" 
                           for k, v in jira_dictionary.items()}
    logger.debug(f"Parsed response - summary length: {len(summary)}, "
                f"updated_stories_map (ID: Title): {stories_with_titles}, "
                f"new_jira_createlist count: {len(new_jira_createlist)}")
    logger.info("Successfully returning response")
    return {
      "jira_dictionary": jira_dictionary,
      "result_text": summary,
      "new_jira_createlist": new_jira_createlist,
    }
    
  except Exception as e:
    logger.error(f"Error in generate() function: {str(e)}", exc_info=True)
    raise

if __name__ == "__main__":
  logger.info("Starting Flask web server")
  port = int(os.environ.get("PORT", 8080))
  logger.info(f"Server will run on port {port}")
  app.run(host='0.0.0.0', port=port, debug=True)