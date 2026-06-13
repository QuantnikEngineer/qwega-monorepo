"""Generate a professional Architecture Document for quantnik-evals in DOCX format."""

from __future__ import annotations

import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ── Colour palette ───────────────────────────────────────────────────

BRAND_DARK  = RGBColor(0x1B, 0x2A, 0x4A)
BRAND_MID   = RGBColor(0x2E, 0x86, 0xAB)
BRAND_LIGHT = RGBColor(0xE8, 0xF1, 0xF5)
PASS_GREEN  = RGBColor(0x27, 0xAE, 0x60)
FAIL_RED    = RGBColor(0xE7, 0x4C, 0x3C)
GREY        = RGBColor(0x7F, 0x8C, 0x8D)
BLACK       = RGBColor(0x33, 0x33, 0x33)


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def _add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = BRAND_DARK
    return heading


def _add_table_with_header(doc, headers, rows, col_widths=None):
    """Add a styled table with header row."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

    # Data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(value)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            if row_idx % 2 == 0:
                _set_cell_shading(cell, "F5F7FA")

    return table


def _add_code_block(doc, code_text):
    """Add a monospaced code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p


def generate_architecture_doc():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BLACK

    # ══════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════════════════════════

    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Quantnik-Evals")
    run.font.size = Pt(36)
    run.font.color.rgb = BRAND_DARK
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Architecture & Design Document")
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND_MID
    run.bold = True

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "A Profile-Driven Evaluation Framework for SDLC AI Agents\n"
        "Version 0.1.0"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = GREY

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Document Date: {datetime.now().strftime('%B %d, %Y')}\n"
        "Classification: INTERNAL"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = GREY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Overview & Objectives",
        "2. High-Level Architecture",
        "3. Core Components",
        "4. Design Patterns",
        "5. Evaluation Methodology",
        "6. Agent Profile System",
        "7. Data Models",
        "8. Configuration & Deployment",
        "9. External Integrations",
        "10. Evaluation Dimensions Taxonomy",
        "11. CLI Reference",
        "12. Extensibility & Customization",
        "13. Current Limitations & Roadmap",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  1. OVERVIEW & OBJECTIVES
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "1. Overview & Objectives", level=1)

    doc.add_paragraph(
        "Quantnik-Evals is a reusable, profile-driven evaluation framework purpose-built "
        "for systematically assessing the quality, reliability, and safety of SDLC AI agents. "
        "It combines LLM-as-Judge assessments with programmatic metrics to provide a "
        "comprehensive quality view of each agent across multiple dimensions."
    )

    _add_styled_heading(doc, "1.1 Key Objectives", level=2)
    objectives = [
        "Standardize evaluation across 11 SDLC AI agents using a common taxonomy",
        "Provide agent-specific evaluation dimensions alongside universal quality metrics",
        "Enable continuous evaluation from production traces (Langfuse)",
        "Support both automated (CLI batch) and interactive (dashboard) evaluation modes",
        "Generate audit-ready reports with traceability back to source traces",
        "Minimize onboarding effort — adding a new agent requires only a profile definition",
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style="List Bullet")

    _add_styled_heading(doc, "1.2 Technology Stack", level=2)
    _add_table_with_header(doc,
        ["Component", "Technology", "Purpose"],
        [
            ["Runtime", "Python ≥3.11", "Framework language"],
            ["LLM Judge", "Google Gemini 2.5 (Vertex AI)", "Structured evaluation scoring"],
            ["Observability", "Langfuse 4.x", "Trace storage, dataset management, score persistence"],
            ["Data Models", "Pydantic 2.x", "Type-safe data contracts"],
            ["Dashboard", "Streamlit", "Real-time evaluation visualization"],
            ["Reporting", "python-docx + matplotlib", "Professional DOCX reports with charts"],
            ["Configuration", "python-dotenv + dataclass", "12-factor environment-driven config"],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  2. HIGH-LEVEL ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "2. High-Level Architecture", level=1)

    doc.add_paragraph(
        "The framework follows a layered architecture with clear separation of concerns:"
    )

    _add_code_block(doc,
        "┌─────────────────────────────────────────────────────────────────┐\n"
        "│                     PRESENTATION LAYER                          │\n"
        "│  ┌──────────┐    ┌───────────────┐    ┌───────────────────┐    │\n"
        "│  │   CLI    │    │   Dashboard   │    │   DOCX Reports    │    │\n"
        "│  │(cli.py)  │    │(dashboard.py) │    │(generate_report)  │    │\n"
        "│  └────┬─────┘    └───────┬───────┘    └─────────┬─────────┘    │\n"
        "├───────┼──────────────────┼──────────────────────┼──────────────┤\n"
        "│       │          ORCHESTRATION LAYER             │              │\n"
        "│       ▼                                          │              │\n"
        "│  ┌──────────────────────────────────┐            │              │\n"
        "│  │         EvalRunner               │            │              │\n"
        "│  │  • Parallel item processing      │            │              │\n"
        "│  │  • Score aggregation & push      │            │              │\n"
        "│  └──────┬───────────┬───────────────┘            │              │\n"
        "├─────────┼───────────┼────────────────────────────┼──────────────┤\n"
        "│         │   EVALUATION ENGINE LAYER              │              │\n"
        "│         ▼           ▼                            │              │\n"
        "│  ┌────────────┐  ┌─────────────────────┐        │              │\n"
        "│  │  LLM Judge │  │Programmatic Evals   │        │              │\n"
        "│  │ (Gemini)   │  │(latency/cost/struct)│        │              │\n"
        "│  └────────────┘  └─────────────────────┘        │              │\n"
        "├─────────────────────────────────────────────────────────────────┤\n"
        "│                    PROFILE LAYER                                │\n"
        "│  ┌──────────────────────────────────────────────────────────┐  │\n"
        "│  │  AgentProfile (base) + 11 registered profiles            │  │\n"
        "│  │  • Dimensions, Judge Prompts, Extraction Logic           │  │\n"
        "│  │  • Auto-merged universal dimensions via __init_subclass__│  │\n"
        "│  └──────────────────────────────────────────────────────────┘  │\n"
        "├─────────────────────────────────────────────────────────────────┤\n"
        "│                    DATA & INTEGRATION LAYER                     │\n"
        "│  ┌──────────────┐  ┌────────────┐  ┌────────────────────┐     │\n"
        "│  │DatasetManager│  │ Langfuse   │  │ Google Vertex AI   │     │\n"
        "│  │              │  │ (traces,   │  │ (Gemini 2.5 judge) │     │\n"
        "│  │              │  │  scores)   │  │                    │     │\n"
        "│  └──────────────┘  └────────────┘  └────────────────────┘     │\n"
        "└─────────────────────────────────────────────────────────────────┘"
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Data Flow: Production traces in Langfuse → Seeded into evaluation datasets → "
        "Evaluated by LLM Judge + Programmatic metrics → Scores pushed back to Langfuse → "
        "Visualized in Dashboard / exported as DOCX reports."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  3. CORE COMPONENTS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "3. Core Components", level=1)

    components = [
        ["AgentProfile\n(agent_profile.py)", "Base class with registry pattern. Defines the contract for agent evaluation: dimensions, judge prompts, trace extraction, programmatic evals. Uses __init_subclass__ to auto-merge universal dimensions."],
        ["LLMJudge\n(llm_judge.py)", "Profile-agnostic LLM-as-judge engine. Renders prompts from profile templates, calls Gemini with structured JSON output, extracts scores. Evaluates dimensions in parallel via ThreadPoolExecutor."],
        ["EvalRunner\n(runner.py)", "Orchestrator that coordinates dataset loading, parallel item evaluation, LLM judge + programmatic scoring, and score persistence to Langfuse."],
        ["DatasetManager\n(dataset_manager.py)", "Manages evaluation datasets: seeds from Langfuse traces, pushes/loads datasets, adds manual items. Uses profile extractors to normalize trace data."],
        ["EvalConfig\n(config.py)", "Dataclass configuration with environment variable defaults. Controls judge model, parallelism, thresholds, and connection settings."],
        ["Common Prompts\n(common_prompts.py)", "Library of 11 universal evaluation prompt templates and dimension groupings (UNIVERSAL_DIMENSIONS, REASONING_DIMENSIONS)."],
        ["CLI\n(cli.py)", "argparse-based interface with subcommands: list-agents, seed, run, trace, add-item. Entry point: `quantnik-eval`."],
        ["Dashboard\n(dashboard.py)", "Streamlit application fetching scores from Langfuse REST API. Renders KPI cards, bar charts, trend lines, and heatmaps."],
    ]

    _add_table_with_header(doc,
        ["Component", "Responsibility"],
        components,
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  4. DESIGN PATTERNS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "4. Design Patterns", level=1)

    patterns = [
        ["Registry Pattern", "@register_profile decorator populates a module-level dict. Profiles self-register on import via side-effect."],
        ["Template Method", "AgentProfile defines abstract extraction/evaluation hooks. Subclasses override specific steps while inheriting the overall flow."],
        ["Strategy Pattern", "LLMJudge is profile-agnostic — evaluation strategy (prompts, template vars) is injected via the profile object."],
        ["Plugin Architecture", "Adding a new agent = creating one .py file with a decorated class. No core code changes required."],
        ["__init_subclass__ Hook", "Automatically merges universal dimensions and common prompts into every profile subclass at class definition time."],
        ["Parallel Execution", "Two-level ThreadPoolExecutor: items (max_workers=4) and dimensions per item (max_dim_workers=8)."],
        ["Adapter Pattern", "CA bundle builder adapts corporate proxy/SSL environments to standard Python SSL context."],
    ]

    _add_table_with_header(doc,
        ["Pattern", "Application"],
        patterns,
    )

    doc.add_paragraph()

    _add_styled_heading(doc, "4.1 Profile Auto-Merge Mechanism", level=2)
    doc.add_paragraph(
        "When any class inherits from AgentProfile, the __init_subclass__ hook automatically:"
    )
    steps = [
        "Appends UNIVERSAL_DIMENSIONS (8 dimensions) to the subclass's dimension list, avoiding duplicates",
        "Merges COMMON_JUDGE_PROMPTS into the subclass's judge_prompts dict (subclass prompts take priority)",
        "This ensures every agent gets safety, quality, and compliance evaluations without boilerplate",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  5. EVALUATION METHODOLOGY
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "5. Evaluation Methodology", level=1)

    _add_styled_heading(doc, "5.1 Dual Evaluation Approach", level=2)

    approaches = [
        ["LLM-as-Judge", "Gemini 2.5 evaluates agent output against structured rubrics. Each dimension has a custom prompt producing a 0.0–1.0 score with JSON reasoning.", "~13 dimensions per agent"],
        ["Programmatic", "Rule-based metrics computed from trace metadata. Deterministic, fast, no LLM cost.", "3–5 dimensions per agent"],
    ]
    _add_table_with_header(doc,
        ["Approach", "Description", "Coverage"],
        approaches,
    )

    doc.add_paragraph()

    _add_styled_heading(doc, "5.2 Evaluation Flow", level=2)

    _add_code_block(doc,
        "1. Dataset Loading\n"
        "   └─ Load items from Langfuse dataset (trace_id + input/output)\n"
        "\n"
        "2. Per-Item Processing (parallel, max_workers=4)\n"
        "   ├─ Fetch full trace from Langfuse (observations)\n"
        "   ├─ Profile.extract_input() → structured input data\n"
        "   ├─ Profile.extract_output() → structured output data\n"
        "   ├─ Profile.extract_token_info() → latency/cost metadata\n"
        "   │\n"
        "   ├─ LLM Judge Evaluation (parallel, max_dim_workers=8)\n"
        "   │   ├─ Build template vars (input_text, output_text, extras)\n"
        "   │   ├─ For each dimension:\n"
        "   │   │   ├─ Render prompt template\n"
        "   │   │   ├─ Call Gemini (JSON response)\n"
        "   │   │   └─ Parse score + reasoning\n"
        "   │   └─ Collect DimensionScore objects\n"
        "   │\n"
        "   ├─ Programmatic Evaluation\n"
        "   │   ├─ Structural validity (schema check)\n"
        "   │   ├─ Token efficiency\n"
        "   │   ├─ Latency scoring\n"
        "   │   ├─ Cost efficiency\n"
        "   │   └─ Agent-specific checks (e.g., severity distribution)\n"
        "   │\n"
        "   └─ Push scores to Langfuse (per-dimension + overall)\n"
        "\n"
        "3. Summary & Reporting\n"
        "   ├─ Aggregate mean/min/max per dimension\n"
        "   ├─ Apply pass threshold (default 0.7)\n"
        "   └─ Log PASS/FAIL summary"
    )

    doc.add_paragraph()

    _add_styled_heading(doc, "5.3 Scoring Conventions", level=2)
    scoring = [
        ["0.90 – 1.00", "Excellent"],
        ["0.70 – 0.89", "Good (PASS)"],
        ["0.50 – 0.69", "Fair"],
        ["0.30 – 0.49", "Poor"],
        ["0.00 – 0.29", "Critical"],
    ]
    _add_table_with_header(doc, ["Score Range", "Rating"], scoring)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  6. AGENT PROFILE SYSTEM
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "6. Agent Profile System", level=1)

    doc.add_paragraph(
        "Each SDLC AI agent is represented by a Profile class that encapsulates all "
        "evaluation logic specific to that agent. The profile system provides:"
    )
    bullets = [
        "Agent-specific evaluation dimensions and judge prompts",
        "Custom trace extraction logic (how to find input/output in Langfuse observations)",
        "Extra template variables for judge prompts (e.g., {framework_type}, {diff})",
        "Optional programmatic evaluators beyond the base set",
        "Output schema declaration for structural validation",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    _add_styled_heading(doc, "6.1 Profile Contract (Template Methods)", level=2)

    methods = [
        ["extract_input(trace, observations)", "Extract structured input from Langfuse trace", "dict[str, Any]"],
        ["extract_output(trace, observations)", "Extract structured output from observations", "dict | None"],
        ["extract_primary_input_text(input_data)", "Primary text for {input_text} template var", "str"],
        ["extract_primary_output_text(output_data)", "Primary text for {output_text} template var", "str"],
        ["format_judge_context(input_data, output_data)", "Extra template variables for prompts", "dict[str, str]"],
        ["run_programmatic_evals(input, output, token_info)", "Agent-specific programmatic checks", "list[DimensionScore]"],
        ["extract_token_info(observations)", "Token count, latency, cost from traces", "dict[str, int]"],
    ]
    _add_table_with_header(doc, ["Method", "Purpose", "Returns"], methods)

    doc.add_paragraph()

    _add_styled_heading(doc, "6.2 Registered Profiles (11 Agents)", level=2)

    profiles = [
        ["cara", "cara_prompt_stream", "Code review (CARA)", "8", "Yes"],
        ["brd_summary", "brd_summary_request", "BRD summarization", "5", "No"],
        ["brd", "brd_chat", "BRD generation chat", "6", "No"],
        ["user_story", "user_story_gen", "User story generation", "5", "Yes"],
        ["userstory_validator", "userstory_validator", "Story validation", "6", "No"],
        ["userstory_to_testcases", "userstory_to_testcases", "Story → Test cases", "6", "No"],
        ["testcases_to_testdata", "testcase_to_testdata_conversion", "TC → Test data", "4", "No"],
        ["testcase_to_scripts", "testcase_to_script_conversion", "TC → Scripts", "5", "Yes"],
        ["sdlc_orchestrator", "sdlc_orchestrator_request", "Intent routing", "5", "No"],
        ["code_assistant", "codeassist", "Code assistance", "7", "No"],
        ["user_manual", "usermanual_generation", "Manual generation", "6", "No"],
    ]
    _add_table_with_header(doc,
        ["Profile Name", "Trace Name", "Purpose", "Domain Dims", "Custom Extractors"],
        profiles,
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Note: Every profile automatically inherits 8 universal dimensions "
        "(hallucination, faithfulness, coherence, conciseness, consistency, toxicity, "
        "data_privacy_compliance, policy_compliance) plus 3 programmatic dimensions "
        "(latency, cost_efficiency, token_efficiency, structural_validity). "
        "Total dimensions per agent: 15–20."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  7. DATA MODELS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "7. Data Models", level=1)

    doc.add_paragraph(
        "All data models use Pydantic v2 BaseModel for type safety and serialization."
    )

    models = [
        ["DimensionScore", "dimension: str, score: float (0–1), reasoning: str, details: dict, evaluator: str", "Single evaluation dimension result"],
        ["EvalResult", "trace_id, agent, run_name, scores: list[DimensionScore], error: str|None", "Complete evaluation for one item. overall_score property = mean of scores"],
        ["DatasetItem", "id, input: dict, expected_output: dict|None, metadata: dict, trace_id: str|None", "Single evaluation dataset item"],
        ["ItemEval", "item_key, is_correct, score, reasoning, metadata", "Per-finding evaluation detail"],
    ]
    _add_table_with_header(doc, ["Model", "Key Fields", "Purpose"], models)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  8. CONFIGURATION & DEPLOYMENT
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "8. Configuration & Deployment", level=1)

    _add_styled_heading(doc, "8.1 Environment Variables", level=2)

    config_vars = [
        ["LANGFUSE_PUBLIC_KEY", "Required", "Langfuse authentication"],
        ["LANGFUSE_SECRET_KEY", "Required", "Langfuse authentication"],
        ["LANGFUSE_HOST", "Required", "Langfuse server URL"],
        ["JUDGE_MODEL", "gemini-2.5-pro", "LLM judge model name"],
        ["JUDGE_USE_VERTEX", "true", "Use Vertex AI (vs API key)"],
        ["VERTEX_PROJECT", "—", "GCP project for Vertex AI"],
        ["VERTEX_LOCATION", "us-central1", "GCP region"],
        ["EVAL_MAX_WORKERS", "4", "Parallel items"],
        ["EVAL_MAX_DIM_WORKERS", "8", "Parallel dimensions per item"],
        ["EVAL_PASS_THRESHOLD", "0.7", "Pass/fail threshold"],
        ["SSL_CERT_FILE", "auto-generated", "CA bundle path (corporate proxy)"],
    ]
    _add_table_with_header(doc, ["Variable", "Default", "Purpose"], config_vars)

    doc.add_paragraph()

    _add_styled_heading(doc, "8.2 Installation", level=2)
    _add_code_block(doc,
        "# Install in development mode\n"
        "pip install -e .\n"
        "\n"
        "# Or with dashboard support\n"
        "pip install -e .[dashboard]\n"
        "\n"
        "# Verify\n"
        "quantnik-eval list-agents"
    )

    _add_styled_heading(doc, "8.3 SSL / Corporate Proxy Handling", level=2)
    doc.add_paragraph(
        "The framework automatically handles self-signed certificates and corporate "
        "proxy environments (e.g., Zscaler). On first run, _ensure_ca_bundle() builds "
        "a combined CA certificate chain from: (1) certifi's Mozilla CA bundle, "
        "(2) the OS certificate store (Windows/macOS), and (3) the Langfuse server's "
        "TLS certificate fetched live. This bundle is cached in a temp directory and "
        "set as SSL_CERT_FILE for the process lifetime."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  9. EXTERNAL INTEGRATIONS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "9. External Integrations", level=1)

    integrations = [
        ["Langfuse", "Trace ingestion, dataset storage, score persistence", "SDK + REST API", "Trace observations, dataset items, numeric scores"],
        ["Google Vertex AI", "LLM-as-Judge model hosting", "google-genai SDK", "Structured JSON evaluation responses"],
        ["Gemini 2.5", "Judge model for dimension scoring", "via Vertex AI", "0.0–1.0 scores with reasoning"],
    ]
    _add_table_with_header(doc,
        ["System", "Role", "Integration Method", "Data Exchanged"],
        integrations,
    )

    doc.add_paragraph()

    _add_styled_heading(doc, "9.1 Langfuse Integration Details", level=2)
    doc.add_paragraph("The framework interacts with Langfuse in three ways:")
    langfuse_ops = [
        "READ: Fetch production traces (observations, input/output, metadata) for seeding datasets",
        "WRITE: Create/update datasets with evaluation items (input + expected_output + metadata)",
        "WRITE: Push evaluation scores back to traces (per-dimension + aggregate, NUMERIC type)",
    ]
    for op in langfuse_ops:
        doc.add_paragraph(op, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Score naming convention: {agent_name}_eval_{dimension} (e.g., cara_eval_hallucination). "
        "This enables filtering and grouping in the Langfuse UI and the Streamlit dashboard."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  10. EVALUATION DIMENSIONS TAXONOMY
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "10. Evaluation Dimensions Taxonomy", level=1)

    doc.add_paragraph(
        "Dimensions are organized into categories. Universal dimensions apply to ALL agents; "
        "domain-specific dimensions are defined per profile."
    )

    _add_styled_heading(doc, "10.1 Universal Dimensions (auto-applied to all agents)", level=2)
    universal = [
        ["hallucination", "Safety & Trust", "LLM Judge", "Detects fabricated/unsupported information"],
        ["faithfulness", "Output Quality", "LLM Judge", "Grounding in source material"],
        ["coherence", "Output Quality", "LLM Judge", "Logical structure and flow"],
        ["conciseness", "Output Quality", "LLM Judge", "Absence of unnecessary verbosity"],
        ["consistency", "Output Quality", "LLM Judge", "Uniform terminology and style"],
        ["toxicity", "Safety & Trust", "LLM Judge", "Harmful/inappropriate content screening"],
        ["data_privacy_compliance", "Safety & Trust", "LLM Judge", "PII/secrets detection"],
        ["policy_compliance", "Compliance", "LLM Judge", "Enterprise standards adherence"],
    ]
    _add_table_with_header(doc, ["Dimension", "Category", "Evaluator", "Description"], universal)

    doc.add_paragraph()

    _add_styled_heading(doc, "10.2 Programmatic Dimensions (auto-applied)", level=2)
    programmatic = [
        ["structural_validity", "Structural", "Schema Check", "Output conforms to declared schema"],
        ["latency", "Performance", "Threshold", "Response time: ≤10s=1.0, ≤30s=0.8, ≤60s=0.6, ≤120s=0.4"],
        ["cost_efficiency", "Performance", "Threshold", "Token usage: ≤5K=1.0, ≤20K=0.8, ≤50K=0.6"],
        ["token_efficiency", "Performance", "Ratio", "Output-to-total token ratio"],
    ]
    _add_table_with_header(doc, ["Dimension", "Category", "Method", "Description"], programmatic)

    doc.add_paragraph()

    _add_styled_heading(doc, "10.3 Domain-Specific Dimensions (examples)", level=2)
    domain = [
        ["CARA (Code Review)", "review_relevance, finding_accuracy, severity_calibration, false_positive_rate, remediation_quality"],
        ["Test Scripts", "script_correctness, framework_compliance, test_coverage_mapping, code_quality, executability"],
        ["SDLC Orchestrator", "intent_routing_accuracy, response_quality, context_retention, flow_orchestration"],
        ["BRD Generation", "reply_quality, requirement_elicitation, brd_document_quality, conversation_flow"],
        ["User Story", "story_quality, acceptance_criteria_quality, coverage, invest_compliance, story_sizing"],
    ]
    _add_table_with_header(doc, ["Agent", "Domain Dimensions"], domain)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  11. CLI REFERENCE
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "11. CLI Reference", level=1)

    doc.add_paragraph("Entry point: quantnik-eval (installed via pip)")

    cli_cmds = [
        ["list-agents", "—", "List all registered agent profiles"],
        ["seed", "--agent, --dataset, --limit, --trace-name, --min-obs", "Seed dataset from Langfuse production traces"],
        ["run", "--agent, --dataset, --concurrency, --dimensions, --max-items, --judge-model, -o", "Run evaluation on a dataset"],
        ["trace", "--agent, <trace_id>, --dimensions, --judge-model", "Evaluate a single trace"],
        ["add-item", "--agent, --dataset, --input-file, --expected-file, --item-id", "Add manual dataset item"],
    ]
    _add_table_with_header(doc, ["Command", "Key Arguments", "Description"], cli_cmds)

    doc.add_paragraph()
    _add_styled_heading(doc, "11.1 Usage Examples", level=2)
    _add_code_block(doc,
        "# List all available agents\n"
        "quantnik-eval list-agents\n"
        "\n"
        "# Seed dataset from production traces\n"
        "quantnik-eval --agent cara seed --limit 50\n"
        "\n"
        "# Run full evaluation\n"
        "quantnik-eval --agent cara run --dataset cara-eval --concurrency 4\n"
        "\n"
        "# Evaluate specific dimensions only\n"
        "quantnik-eval --agent cara run --dimensions hallucination,toxicity\n"
        "\n"
        "# Evaluate a single trace\n"
        "quantnik-eval --agent cara trace abc123def456"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  12. EXTENSIBILITY & CUSTOMIZATION
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "12. Extensibility & Customization", level=1)

    _add_styled_heading(doc, "12.1 Adding a New Agent (Minimal Profile)", level=2)
    doc.add_paragraph(
        "Adding evaluation for a new agent requires creating a single Python file "
        "in quantnik_evals/profiles/ with a decorated class:"
    )
    _add_code_block(doc,
        '@register_profile\n'
        'class MyNewAgentProfile(AgentProfile):\n'
        '    name = "my_new_agent"\n'
        '    description = "Evaluates My New Agent"\n'
        '    trace_name = "my_agent_trace_name"  # Langfuse trace filter\n'
        '    default_dataset = "my-agent-eval"\n'
        '\n'
        '    dimensions = [\n'
        '        "accuracy",\n'
        '        "relevance",\n'
        '    ]\n'
        '\n'
        '    judge_prompts = {\n'
        '        "accuracy": """Evaluate the accuracy...\\n{output_text}""",\n'
        '        "relevance": """Assess relevance...\\n{input_text}\\n{output_text}""",\n'
        '    }\n'
        '\n'
        '    # 8 universal dimensions + prompts auto-merged!\n'
        '    # No boilerplate needed for hallucination, toxicity, etc.'
    )

    doc.add_paragraph()

    _add_styled_heading(doc, "12.2 Customization Points", level=2)
    customizations = [
        ["Evaluation Dimensions", "Add to profile's `dimensions` list + corresponding `judge_prompts` entry"],
        ["Trace Extraction", "Override `extract_input()` / `extract_output()` for non-standard trace structures"],
        ["Judge Template Variables", "Override `format_judge_context()` to provide agent-specific variables (e.g., {framework_type})"],
        ["Programmatic Checks", "Override `run_programmatic_evals()` to add custom heuristics"],
        ["Output Schema", "Set `output_schema` dict to enable automatic structural validation"],
        ["Pass Threshold", "Set via EVAL_PASS_THRESHOLD env var or CLI"],
        ["Judge Model", "Set via --judge-model CLI arg or JUDGE_MODEL env var"],
        ["Parallelism", "Set via --concurrency CLI arg or EVAL_MAX_WORKERS env var"],
    ]
    _add_table_with_header(doc, ["What to Customize", "How"], customizations)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  13. CURRENT LIMITATIONS & ROADMAP
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "13. Current Limitations & Roadmap", level=1)

    _add_styled_heading(doc, "13.1 Current Scope (Layer 1 Only)", level=2)
    doc.add_paragraph(
        "The framework currently supports Agent-Level Evaluation only. "
        "Each agent is evaluated in isolation against its own dimensions."
    )

    limitations = [
        ["Single-agent scope", "No evaluation of handoffs between agents or end-to-end pipelines"],
        ["No weighted scoring", "Overall score is simple average; no dimension weighting"],
        ["No reference comparison", "Cannot compare against golden reference outputs"],
        ["No result caching", "Re-running evaluates from scratch (no skip-if-already-scored)"],
        ["Hardcoded thresholds", "Latency/cost scoring tiers are not configurable per agent"],
        ["No prompt versioning", "Cannot A/B test prompt template changes"],
    ]
    _add_table_with_header(doc, ["Limitation", "Impact"], limitations)

    doc.add_paragraph()

    _add_styled_heading(doc, "13.2 Planned Layers", level=2)

    roadmap = [
        ["Layer 2: Transition Evaluation", "Planned", "Evaluate handoffs between agents — schema compatibility, semantic fidelity, data loss detection"],
        ["Layer 3: End-to-End Evaluation", "Planned", "Evaluate complete multi-agent pipelines (e.g., BRD → Stories → TCs → Scripts) with quality degradation tracking"],
        ["Weighted Scoring", "Planned", "Allow per-dimension weights (e.g., hallucination=3x, conciseness=0.5x)"],
        ["YAML-driven Dimensions", "Planned", "Non-developers define new evaluation dimensions via configuration files"],
        ["Comparative Evaluation", "Planned", "Side-by-side comparison of Agent v1 vs v2 on same inputs"],
        ["Multi-Judge Consensus", "Future", "Run multiple judge models for confidence scoring"],
    ]
    _add_table_with_header(doc, ["Feature", "Status", "Description"], roadmap)

    doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Generated by quantnik-evals \u2022 {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

    # ── Save ──────────────────────────────────────────────────────────
    output_path = os.path.join(
        os.path.dirname(__file__), "quantnik_evals_architecture.docx"
    )
    doc.save(output_path)
    print(f"Architecture document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_architecture_doc()
