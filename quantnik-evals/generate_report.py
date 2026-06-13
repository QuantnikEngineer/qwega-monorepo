"""Generate a professional DOCX evaluation report for an agent."""

from __future__ import annotations

import io
import os
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# ── Evaluation Data ──────────────────────────────────────────────────

AGENT_NAME = "testcases_to_testdata"
AGENT_DESCRIPTION = (
    "Converts structured test cases into realistic, executable test data sets. "
    "The agent interprets test case specifications and generates corresponding "
    "data values including boundary conditions, valid/invalid inputs, and edge cases."
)
EVAL_DATE = "May 15, 2026"
DATASET = "tc-to-testdata-eval"
ITEMS_EVALUATED = 85
PASS_THRESHOLD = 0.70
JUDGE_MODEL = "gemini-2.5-flash (Google Vertex AI)"

# Combined results from both runs
DIMENSIONS = {
    # Domain-specific (Run 1)
    "boundary_values":        {"score": 0.06, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    "data_completeness":      {"score": 0.27, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    "data_validity":          {"score": 0.28, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    "data_variety":           {"score": 0.18, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    "structural_validity":    {"score": 0.06, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    # Common LLM Judge (Run 2)
    "coherence":              {"score": 0.27, "category": "Output Quality",  "evaluator": "LLM Judge"},
    "conciseness":            {"score": 0.28, "category": "Output Quality",  "evaluator": "LLM Judge"},
    "consistency":            {"score": 0.97, "category": "Output Quality",  "evaluator": "LLM Judge"},
    "faithfulness":           {"score": 0.29, "category": "Output Quality",  "evaluator": "LLM Judge"},
    "hallucination":          {"score": 0.79, "category": "Safety & Trust",  "evaluator": "LLM Judge"},
    "toxicity":               {"score": 1.00, "category": "Safety & Trust",  "evaluator": "LLM Judge"},
    "data_privacy_compliance":{"score": 0.88, "category": "Safety & Trust",  "evaluator": "LLM Judge"},
    "policy_compliance":      {"score": 0.34, "category": "Compliance",      "evaluator": "LLM Judge"},
    # Programmatic (Run 2)
    "latency":                {"score": 1.00, "category": "Performance",     "evaluator": "Programmatic"},
    "cost_efficiency":        {"score": 1.00, "category": "Performance",     "evaluator": "Programmatic"},
}

# ── Colour palette ───────────────────────────────────────────────────

BRAND_DARK  = RGBColor(0x1B, 0x2A, 0x4A)   # dark navy
BRAND_MID   = RGBColor(0x2E, 0x86, 0xAB)   # teal
BRAND_LIGHT = RGBColor(0xE8, 0xF1, 0xF5)   # light blue-grey
PASS_GREEN  = RGBColor(0x27, 0xAE, 0x60)
FAIL_RED    = RGBColor(0xE7, 0x4C, 0x3C)
GREY        = RGBColor(0x7F, 0x8C, 0x8D)


def _set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def _add_styled_heading(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = BRAND_DARK
    return heading


def _create_bar_chart(dimensions: dict, filename: str):
    """Create a horizontal bar chart of dimension scores."""
    sorted_dims = sorted(dimensions.items(), key=lambda x: x[1]["score"], reverse=True)
    names = [d[0].replace("_", " ").title() for d in sorted_dims]
    scores = [d[1]["score"] for d in sorted_dims]
    colors = ["#27AE60" if s >= PASS_THRESHOLD else "#E74C3C" for s in scores]

    fig, ax = plt.subplots(figsize=(8, 6))
    bars = ax.barh(names, scores, color=colors, height=0.6, edgecolor="white", linewidth=0.5)
    ax.axvline(x=PASS_THRESHOLD, color="#2E86AB", linestyle="--", linewidth=1.5, label=f"Pass Threshold ({PASS_THRESHOLD})")
    ax.set_xlim(0, 1.05)
    ax.set_xlabel("Score", fontsize=11, fontweight="bold")
    ax.set_title("Dimension Scores Overview", fontsize=14, fontweight="bold", color="#1B2A4A", pad=15)
    ax.legend(loc="lower right", fontsize=9)
    ax.invert_yaxis()

    for bar, score in zip(bars, scores):
        ax.text(score + 0.02, bar.get_y() + bar.get_height()/2, f"{score:.2f}",
                va="center", fontsize=9, fontweight="bold")

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    plt.savefig(filename, dpi=200, bbox_inches="tight")
    plt.close()


def _create_category_radar(dimensions: dict, filename: str):
    """Create a radar/spider chart by category."""
    categories: dict[str, list[float]] = {}
    for dim, info in dimensions.items():
        categories.setdefault(info["category"], []).append(info["score"])
    cat_means = {cat: sum(s)/len(s) for cat, s in categories.items()}

    labels = list(cat_means.keys())
    values = list(cat_means.values())
    N = len(labels)

    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    values += values[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color="#2E86AB", alpha=0.25)
    ax.plot(angles, values, color="#2E86AB", linewidth=2)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=10, fontweight="bold")
    ax.set_ylim(0, 1.0)
    ax.set_title("Category Performance", fontsize=14, fontweight="bold", color="#1B2A4A", pad=20)

    # Add threshold circle
    threshold_values = [PASS_THRESHOLD] * (N + 1)
    ax.plot(angles, threshold_values, color="#E74C3C", linestyle="--", linewidth=1, alpha=0.7)

    plt.tight_layout()
    plt.savefig(filename, dpi=200, bbox_inches="tight")
    plt.close()


def _create_pass_fail_pie(dimensions: dict, filename: str):
    """Create a pass/fail pie chart."""
    passing = sum(1 for d in dimensions.values() if d["score"] >= PASS_THRESHOLD)
    failing = len(dimensions) - passing

    fig, ax = plt.subplots(figsize=(5, 5))
    wedges, texts, autotexts = ax.pie(
        [passing, failing],
        labels=["Pass", "Fail"],
        colors=["#27AE60", "#E74C3C"],
        autopct="%1.0f%%",
        startangle=90,
        textprops={"fontsize": 13, "fontweight": "bold"},
        wedgeprops={"edgecolor": "white", "linewidth": 2},
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontsize(14)
    ax.set_title(f"Pass/Fail Breakdown ({passing}/{len(dimensions)} dimensions)",
                 fontsize=13, fontweight="bold", color="#1B2A4A", pad=15)
    plt.tight_layout()
    plt.savefig(filename, dpi=200, bbox_inches="tight")
    plt.close()


def generate_report():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ──────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ══════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Agent Evaluation Report")
    run.font.size = Pt(32)
    run.font.color.rgb = BRAND_DARK
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Agent: {AGENT_NAME.replace('_', ' ').title()}")
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_MID
    run.bold = True

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Evaluation Date: {EVAL_DATE}\nFramework: quantnik-evals v0.1.0\nJudge Model: {JUDGE_MODEL}")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("─" * 60)
    run.font.color.rgb = BRAND_MID

    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = classification.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
    run.font.size = Pt(10)
    run.font.color.rgb = FAIL_RED
    run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS (placeholder)
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Evaluation Methodology",
        "3. Results Overview",
        "4. Detailed Dimension Analysis",
        "5. Category Performance",
        "6. Strengths & Weaknesses",
        "7. Recommendations",
        "8. Appendix",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "1. Executive Summary", level=1)

    all_scores = [d["score"] for d in DIMENSIONS.values()]
    overall = sum(all_scores) / len(all_scores)
    passing = sum(1 for s in all_scores if s >= PASS_THRESHOLD)
    failing = len(all_scores) - passing
    verdict = "PASS" if overall >= PASS_THRESHOLD else "NEEDS IMPROVEMENT"

    doc.add_paragraph(
        f"This report presents the evaluation results for the '{AGENT_NAME}' AI agent, "
        f"assessed across {len(DIMENSIONS)} quality dimensions using {ITEMS_EVALUATED} "
        f"production traces from the '{DATASET}' dataset."
    )

    doc.add_paragraph(AGENT_DESCRIPTION)

    # KPI Table
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.style = "Table Grid"

    headers = ["Overall Score", "Verdict", "Dimensions Passing", "Items Evaluated"]
    values_row = [f"{overall:.2f}", verdict, f"{passing}/{len(DIMENSIONS)}", str(ITEMS_EVALUATED)]
    colors_row = [
        BRAND_MID.hex() if overall >= PASS_THRESHOLD else "E74C3C",
        "27AE60" if verdict == "PASS" else "E74C3C",
        "2E86AB",
        "2E86AB",
    ]

    for i, (h, v) in enumerate(zip(headers, values_row)):
        cell = kpi_table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            run.bold = True

        cell2 = kpi_table.cell(1, i)
        cell2.text = v
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.size = Pt(16)
            run.bold = True
            if i == 0:
                run.font.color.rgb = PASS_GREEN if overall >= PASS_THRESHOLD else FAIL_RED
            elif i == 1:
                run.font.color.rgb = PASS_GREEN if verdict == "PASS" else FAIL_RED
            else:
                run.font.color.rgb = BRAND_MID

    doc.add_paragraph()

    doc.add_paragraph(
        f"The agent achieved an overall score of {overall:.2f} against a pass threshold of "
        f"{PASS_THRESHOLD:.2f}. Out of {len(DIMENSIONS)} evaluated dimensions, "
        f"{passing} passed and {failing} failed. The primary areas of concern are in "
        f"domain-specific quality metrics, particularly boundary value generation, "
        f"structural validity, and data variety."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  2. EVALUATION METHODOLOGY
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "2. Evaluation Methodology", level=1)

    _add_styled_heading(doc, "2.1 Framework", level=2)
    doc.add_paragraph(
        "Evaluations were conducted using the quantnik-evals framework (v0.1.0), a reusable "
        "LLM agent evaluation system that combines LLM-as-judge assessments with "
        "programmatic metrics. All evaluation scores are stored in Langfuse for "
        "traceability and audit purposes."
    )

    _add_styled_heading(doc, "2.2 Evaluation Approaches", level=2)

    approach_table = doc.add_table(rows=3, cols=3)
    approach_table.style = "Table Grid"
    approach_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    approach_headers = ["Approach", "Dimensions", "Description"]
    for i, h in enumerate(approach_headers):
        cell = approach_table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

    llm_dims = [d for d, info in DIMENSIONS.items() if info["evaluator"] == "LLM Judge"]
    prog_dims = [d for d, info in DIMENSIONS.items() if info["evaluator"] == "Programmatic"]

    approaches = [
        ("LLM-as-Judge", str(len(llm_dims)),
         f"Gemini 2.5 Flash evaluates agent output against structured rubrics. "
         f"Each dimension has a custom prompt template producing a 0-1 score with reasoning."),
        ("Programmatic", str(len(prog_dims)),
         f"Rule-based metrics computed from trace metadata: latency (response time), "
         f"cost efficiency (token usage)."),
    ]
    for row_idx, (approach, count, desc) in enumerate(approaches, 1):
        approach_table.cell(row_idx, 0).text = approach
        approach_table.cell(row_idx, 1).text = count
        approach_table.cell(row_idx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        approach_table.cell(row_idx, 2).text = desc

    doc.add_paragraph()

    _add_styled_heading(doc, "2.3 Dataset", level=2)
    doc.add_paragraph(
        f"The evaluation dataset '{DATASET}' contains {ITEMS_EVALUATED} items seeded from "
        f"production Langfuse traces of type 'testcase_to_testdata_conversion'. Each item "
        f"includes the original test case input and the agent's generated test data output."
    )

    _add_styled_heading(doc, "2.4 Scoring", level=2)
    doc.add_paragraph(
        f"All scores are normalized to a 0.0–1.0 scale. A dimension is considered passing "
        f"if its mean score across all items is ≥ {PASS_THRESHOLD:.2f}. The overall score "
        f"is the arithmetic mean of all dimension scores."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  3. RESULTS OVERVIEW
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "3. Results Overview", level=1)

    # Bar chart
    bar_chart_path = os.path.join(os.path.dirname(__file__), "_report_bar.png")
    _create_bar_chart(DIMENSIONS, bar_chart_path)
    doc.add_picture(bar_chart_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Full results table
    _add_styled_heading(doc, "3.1 Dimension Scores", level=2)

    sorted_dims = sorted(DIMENSIONS.items(), key=lambda x: x[1]["score"], reverse=True)
    results_table = doc.add_table(rows=len(sorted_dims) + 1, cols=5)
    results_table.style = "Table Grid"
    results_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_headers = ["Dimension", "Category", "Evaluator", "Score", "Status"]
    for i, h in enumerate(col_headers):
        cell = results_table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

    for row_idx, (dim, info) in enumerate(sorted_dims, 1):
        score = info["score"]
        status = "PASS" if score >= PASS_THRESHOLD else "FAIL"

        results_table.cell(row_idx, 0).text = dim.replace("_", " ").title()
        results_table.cell(row_idx, 1).text = info["category"]
        results_table.cell(row_idx, 2).text = info["evaluator"]

        score_cell = results_table.cell(row_idx, 3)
        score_cell.text = f"{score:.2f}"
        score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in score_cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = PASS_GREEN if score >= PASS_THRESHOLD else FAIL_RED

        status_cell = results_table.cell(row_idx, 4)
        status_cell.text = status
        status_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(status_cell, "27AE60" if status == "PASS" else "E74C3C")
        for run in status_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

        # Alternate row shading
        if row_idx % 2 == 0:
            for col in range(3):
                _set_cell_shading(results_table.cell(row_idx, col), "F5F7FA")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  4. DETAILED DIMENSION ANALYSIS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "4. Detailed Dimension Analysis", level=1)

    dim_descriptions = {
        "boundary_values": (
            "Measures whether the generated test data includes appropriate boundary and edge case values "
            "(e.g., min/max values, empty strings, zero, negative numbers, overflow values).",
            "The agent scored critically low (0.06), indicating it rarely generates boundary conditions. "
            "Most outputs contain only 'happy path' data without stress-testing edge cases."
        ),
        "data_completeness": (
            "Evaluates whether all required data fields specified in the test case are present in the output.",
            "At 0.27, the agent frequently omits required fields or leaves them with placeholder values. "
            "Coverage of the full test case specification is inconsistent."
        ),
        "data_validity": (
            "Checks whether generated data values conform to expected types, formats, and constraints.",
            "Scoring 0.28, many generated values don't match the expected data types or format constraints "
            "defined in the test cases. Type mismatches and format violations are common."
        ),
        "data_variety": (
            "Measures the diversity of generated test data across different scenarios and combinations.",
            "At 0.18, the agent tends to produce repetitive, homogeneous data. There is insufficient "
            "variation across test data sets to adequately cover different scenarios."
        ),
        "structural_validity": (
            "Verifies the output follows the expected JSON/data structure format.",
            "Critically low at 0.06. The agent frequently produces malformed or incorrectly structured "
            "outputs that don't conform to the expected schema."
        ),
        "coherence": (
            "Evaluates logical consistency and meaningful relationships within the generated data.",
            "At 0.27, generated data often contains internal contradictions or logically inconsistent "
            "values that wouldn't make sense in a real-world scenario."
        ),
        "conciseness": (
            "Measures whether the output is appropriately sized without unnecessary verbosity.",
            "Scoring 0.28, outputs tend to include excessive filler content, redundant fields, "
            "or overly verbose descriptions rather than focused test data."
        ),
        "consistency": (
            "Checks whether the agent produces similar quality outputs for similar inputs.",
            "Excellent score of 0.97. The agent is highly reliable and predictable in its output patterns, "
            "showing minimal variance across similar test case types."
        ),
        "faithfulness": (
            "Evaluates whether the generated data accurately reflects the test case specifications.",
            "At 0.29, the agent often deviates from the original test case requirements, generating "
            "data that doesn't align with the specified scenarios and conditions."
        ),
        "hallucination": (
            "Detects fabricated information not grounded in the input test case.",
            "Good score of 0.79. The agent generally avoids introducing entirely fabricated data, "
            "though some manufactured details appear in complex scenarios."
        ),
        "toxicity": (
            "Screens for harmful, offensive, or inappropriate content in outputs.",
            "Perfect score of 1.00. No toxic or inappropriate content detected in any outputs."
        ),
        "data_privacy_compliance": (
            "Checks that generated data doesn't contain or expose real PII/sensitive information.",
            "Strong score of 0.88. The agent mostly generates synthetic data without real PII, "
            "though a few edge cases showed patterns resembling actual personal data."
        ),
        "policy_compliance": (
            "Verifies adherence to organizational data handling and output format policies.",
            "At 0.34, many outputs don't fully conform to the expected organizational standards "
            "for test data generation, including naming conventions and format requirements."
        ),
        "latency": (
            "Measures response time efficiency (≤10s = 1.0, ≤30s = 0.8, ≤60s = 0.6).",
            "Perfect score of 1.00. All traces completed within the optimal latency window (≤10s)."
        ),
        "cost_efficiency": (
            "Evaluates token consumption efficiency (≤5K tokens = 1.0, ≤20K = 0.8).",
            "Perfect score of 1.00. Token usage is highly efficient across all evaluated traces."
        ),
    }

    for dim, info in sorted_dims:
        score = info["score"]
        status = "PASS" if score >= PASS_THRESHOLD else "FAIL"
        desc, analysis = dim_descriptions.get(dim, ("No description available.", "No analysis available."))

        heading = doc.add_heading(level=3)
        run = heading.add_run(f"{dim.replace('_', ' ').title()}  ")
        run.font.color.rgb = BRAND_DARK

        badge_run = heading.add_run(f"  {score:.2f} — {status}  ")
        badge_run.font.size = Pt(11)
        badge_run.font.color.rgb = PASS_GREEN if status == "PASS" else FAIL_RED
        badge_run.bold = True

        doc.add_paragraph(desc).runs[0].font.italic = True
        doc.add_paragraph(analysis)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  5. CATEGORY PERFORMANCE
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "5. Category Performance", level=1)

    # Radar chart
    radar_path = os.path.join(os.path.dirname(__file__), "_report_radar.png")
    _create_category_radar(DIMENSIONS, radar_path)
    doc.add_picture(radar_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Category summary table
    categories: dict[str, list[float]] = {}
    for dim, info in DIMENSIONS.items():
        categories.setdefault(info["category"], []).append(info["score"])

    cat_table = doc.add_table(rows=len(categories) + 1, cols=4)
    cat_table.style = "Table Grid"
    cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["Category", "Avg Score", "Dimensions", "Status"]):
        cell = cat_table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

    for row_idx, (cat, scores) in enumerate(sorted(categories.items()), 1):
        mean = sum(scores) / len(scores)
        status = "PASS" if mean >= PASS_THRESHOLD else "FAIL"
        cat_table.cell(row_idx, 0).text = cat
        score_cell = cat_table.cell(row_idx, 1)
        score_cell.text = f"{mean:.2f}"
        score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in score_cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = PASS_GREEN if mean >= PASS_THRESHOLD else FAIL_RED
        cat_table.cell(row_idx, 2).text = str(len(scores))
        cat_table.cell(row_idx, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_cell = cat_table.cell(row_idx, 3)
        status_cell.text = status
        status_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(status_cell, "27AE60" if status == "PASS" else "E74C3C")
        for run in status_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  6. STRENGTHS & WEAKNESSES
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "6. Strengths & Weaknesses", level=1)

    # Pie chart
    pie_path = os.path.join(os.path.dirname(__file__), "_report_pie.png")
    _create_pass_fail_pie(DIMENSIONS, pie_path)
    doc.add_picture(pie_path, width=Inches(4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _add_styled_heading(doc, "6.1 Key Strengths", level=2)
    strengths = [
        ("Consistency (0.97)", "The agent produces highly predictable and reliable outputs across similar inputs, indicating stable underlying logic."),
        ("Safety & Trust (Avg 0.89)", "Excellent performance in toxicity (1.00), hallucination prevention (0.79), and data privacy compliance (0.88)."),
        ("Performance Efficiency (1.00)", "Both latency and cost efficiency are perfect, meaning the agent is fast and token-efficient."),
        ("Zero Evaluation Errors", f"All {ITEMS_EVALUATED} items were evaluated without any processing errors, demonstrating robust trace quality."),
    ]
    for title, desc in strengths:
        p = doc.add_paragraph()
        run = p.add_run(f"✓ {title}: ")
        run.bold = True
        run.font.color.rgb = PASS_GREEN
        p.add_run(desc)

    _add_styled_heading(doc, "6.2 Critical Weaknesses", level=2)
    weaknesses = [
        ("Structural Validity (0.06)", "The most critical issue — outputs frequently don't conform to the expected schema, making them unusable without post-processing."),
        ("Boundary Values (0.06)", "Near-total failure to generate edge case data. The agent only produces 'happy path' values, missing the core purpose of test data generation."),
        ("Data Variety (0.18)", "Outputs are repetitive with insufficient diversity across generated test data sets."),
        ("Data Completeness (0.27)", "Required fields are frequently missing from the generated test data."),
        ("Faithfulness (0.29)", "Generated data often deviates significantly from the original test case specifications."),
    ]
    for title, desc in weaknesses:
        p = doc.add_paragraph()
        run = p.add_run(f"✗ {title}: ")
        run.bold = True
        run.font.color.rgb = FAIL_RED
        p.add_run(desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  7. RECOMMENDATIONS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "7. Recommendations", level=1)

    doc.add_paragraph(
        "Based on the evaluation results, the following improvements are recommended "
        "in order of priority:"
    )

    recommendations = [
        ("P0 — Fix Output Structure",
         "The agent's output format must be corrected to match the expected JSON schema. "
         "Consider adding structured output constraints or Pydantic model validation to the "
         "agent's prompt and post-processing pipeline. This is the highest-priority fix as "
         "structurally invalid outputs are completely unusable."),
        ("P0 — Enhance Boundary Value Generation",
         "Add explicit instructions in the agent's system prompt to always include boundary "
         "conditions: min/max values, empty inputs, null values, negative numbers, overflow "
         "values, special characters, and format extremes. Consider providing few-shot examples "
         "of good boundary value test data."),
        ("P1 — Improve Data Variety",
         "Adjust the generation strategy to ensure diverse data across test items. Use techniques "
         "like temperature variation, explicit diversity instructions, or combinatorial generation "
         "approaches to avoid repetitive outputs."),
        ("P1 — Strengthen Faithfulness to Specs",
         "Implement a validation step that cross-references generated data against the original "
         "test case specification. Consider a chain-of-thought approach where the agent first "
         "extracts key requirements, then generates data against each requirement."),
        ("P2 — Improve Coherence & Conciseness",
         "Review the agent's prompt to reduce verbose outputs and ensure generated data values "
         "have logical relationships (e.g., dates in order, addresses in valid format)."),
        ("P2 — Policy Compliance",
         "Document and enforce organizational naming conventions, format standards, and test data "
         "policies. Add these as explicit constraints in the agent's system prompt."),
    ]

    for title, desc in recommendations:
        _add_styled_heading(doc, title, level=3)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  8. APPENDIX
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "8. Appendix", level=1)

    _add_styled_heading(doc, "A. Evaluation Configuration", level=2)
    config_table = doc.add_table(rows=8, cols=2)
    config_table.style = "Table Grid"
    config_items = [
        ("Framework", "quantnik-evals v0.1.0"),
        ("Judge Model", JUDGE_MODEL),
        ("Judge Temperature", "0.0"),
        ("Dataset", DATASET),
        ("Items Evaluated", str(ITEMS_EVALUATED)),
        ("Pass Threshold", str(PASS_THRESHOLD)),
        ("Evaluation Date", EVAL_DATE),
        ("Trace Source", "Langfuse (production)"),
    ]
    for i, (k, v) in enumerate(config_items):
        key_cell = config_table.cell(i, 0)
        key_cell.text = k
        for run in key_cell.paragraphs[0].runs:
            run.bold = True
        _set_cell_shading(key_cell, "F5F7FA")
        config_table.cell(i, 1).text = v

    doc.add_paragraph()

    _add_styled_heading(doc, "B. Scoring Scale", level=2)
    doc.add_paragraph(
        "All dimension scores use a 0.0–1.0 continuous scale:\n"
        "• 0.90–1.00: Excellent\n"
        "• 0.70–0.89: Good (Pass)\n"
        "• 0.50–0.69: Fair\n"
        "• 0.30–0.49: Poor\n"
        "• 0.00–0.29: Critical"
    )

    _add_styled_heading(doc, "C. Disclaimer", level=2)
    disclaimer = doc.add_paragraph(
        "This evaluation was performed using automated LLM-as-judge techniques and "
        "programmatic metrics. While the evaluation framework aims for objectivity, "
        "LLM-based assessments may contain inherent biases. Results should be interpreted "
        "in conjunction with human review and domain expertise. Scores reflect the quality "
        "of agent outputs at the time of evaluation and may vary with different input "
        "distributions or updated agent versions."
    )
    disclaimer.runs[0].font.italic = True
    disclaimer.runs[0].font.size = Pt(10)
    disclaimer.runs[0].font.color.rgb = GREY

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Generated by quantnik-evals • {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

    # ── Save ──────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(__file__), f"{AGENT_NAME}_evaluation_report.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

    # Clean up temp chart images
    for path in [bar_chart_path, radar_path, pie_path]:
        if os.path.exists(path):
            os.remove(path)

    return output_path


if __name__ == "__main__":
    generate_report()
