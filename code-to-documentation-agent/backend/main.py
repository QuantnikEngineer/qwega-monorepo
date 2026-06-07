import base64 # For encoding binary file content as base64 for JSON responses
import io
import zipfile
from urllib.parse import urlparse

import httpx
try:
    from docx import Document
except ImportError:
    Document = None
from fastapi import FastAPI, HTTPException, Form, File, UploadFile, Body
from fastapi.middleware.cors import CORSMiddleware
from models import (
    AgentPromptsRequest, AgentPromptsResponse, AgentResponse, ExecuteRequest, 
    SharePointDownloadRequest, SharePointUploadRequest, SharePointUploadResponse, 
    SharePointDownloadResponse, SharePointFileMetadata, SharePointListResponse, 
    SharePointFileItem, 
    GenerateDocumentationRequest, GenerateDocumentationResponse,
    InitSecretsRequest,
    # UploadToConfluenceRequest, UploadToConfluenceResponse,
    # ConfluencePagesResponse, ConfluencePageContentResponse,
    RepoListResponse
)
from secrets_manager import get_secret, init_secrets_context
from fastapi.responses import JSONResponse
from semantic_intent import classify_intent
from run_workflow import execute_workflow, get_agent
from autogen_agentchat.agents import AssistantAgent
from autogen_agentchat.messages import TextMessage
from autogen_core import CancellationToken

from typing import Optional, List
import json
import re
from litellm_client import LiteLLMRunner
# from autogen_ext.tools.mcp import StdioServerParams, mcp_server_tools
import os
from sharepoint_client import SharePointClient, SharePointConfig
# from confluence_client import ConfluenceClient, ConfluenceConfig
import logging
from repo_service import repo_service
from doc_generation_agent import doc_agent
from llm_config import fetch_llm_config


app = FastAPI()

# Simple in-memory store mapping session_id (or user_id) -> last epic/user story text
SESSION_USER_STORIES: dict = {}

# Server-side storage for project_id and auth_provider (initialized from JWT)
APP_CONFIG = {
    "project_id": "",
    "auth_provider": "azure",
}

def _require_env_base_url(env_var: str, description: str) -> str:
    """Ensure mission-critical backend base URL is passed via environment."""
    value = os.getenv(env_var)
    if not value:
        raise RuntimeError(f"{env_var} must be set to reach the {description} service")
    return value.rstrip("/")

# Base URL for the BuildAI backend APIs (used for RAG, agents, connectors, etc.)
BUILDAI_BACKEND_API_URL = _require_env_base_url(
    "BUILDAI_BACKEND_API_URL", "BuildAI backend API"
)

def get_effective_site_name(config: SharePointConfig, logger_instance=None) -> tuple:
    """
    Extract the effective site name and document library from SharePoint configuration.
    
    Configuration fields:
    - site_url: Full URL to SharePoint site (e.g., https://tenant.sharepoint.com/sites/MySite)
    - sharepoint_folder (document_library): Can be in formats:
        - "LibraryName" (e.g., "WEGA")
        - "SiteName/LibraryName" (e.g., "aikidowega/WEGA")
    
    Returns:
        tuple: (site_name, document_library_name)
    """
    log = logger_instance or logging.getLogger(__name__)
    
    # Get the document_library/sharepoint_folder from config
    document_library_config = config.sharepoint_folder or ""
    
    # Extract site name from site_url (this is the authoritative source)
    site_name = None
    if config.site_url and '/sites/' in config.site_url:
        try:
            parsed = urlparse(config.site_url)
            path_parts = parsed.path.strip('/').split('/')
            if len(path_parts) >= 2 and path_parts[0].lower() == 'sites':
                site_name = path_parts[1]
                log.info(f"Extracted site name '{site_name}' from site_url: {config.site_url}")
        except Exception as e:
            log.error(f"Failed to parse site_url: {e}")
    
    # Parse document_library_config - it might be "SiteName/LibraryName" or just "LibraryName"
    document_library = document_library_config
    if "/" in document_library_config:
        # Format is "SiteName/LibraryName" - extract just the library name
        config_site_part, library_part = document_library_config.split("/", 1)
        document_library = library_part
        log.info(f"Parsed document_library config: site_part='{config_site_part}', library='{library_part}'")
        
        # If we couldn't extract site from URL, use the one from document_library config
        if not site_name:
            site_name = config_site_part
            log.warning(f"Using site name from document_library config: '{site_name}'")
    
    if not site_name:
        log.error("Could not extract site name from site_url or document_library config. Check SharePoint configuration.")
        site_name = document_library
        document_library = ""
    
    log.info(f"Effective SharePoint config: site='{site_name}', document_library='{document_library}'")
    return site_name, document_library

@app.get("/llm-config")
async def get_llm_config():
    """Return PROVIDERS and LLM_MODELS loaded from GitHub using secrets."""
    try:
        print(f"[DEBUG] GET /llm-config endpoint called")
        providers, models = await fetch_llm_config()
        return JSONResponse(content={"providers": providers, "models": models})
    except Exception as e:
        return JSONResponse(content={"providers": [], "models": [], "error": str(e)}, status_code=500)

@app.post("/api/init-secrets")
async def init_secrets(request: InitSecretsRequest):
  """Initialize project_id and auth_provider from JWT token values"""
  try:
      project_id = request.project_id
      auth_provider = request.auth_provider

      # Initialize the global context in secrets_manager with JWT values
      init_secrets_context(project_id=project_id, cloud_provider=auth_provider)

      # Store in server memory for backward compatibility
      APP_CONFIG["project_id"] = project_id
      APP_CONFIG["auth_provider"] = auth_provider

      logger.info(
          f"Secrets initialized from JWT - project_id: {project_id}, auth_provider: {auth_provider}"
      )

      return {
          "success": True,
          "project_id": project_id,
          "auth_provider": auth_provider,
      }
  except Exception as e:
      logger.error(f"Failed to initialize secrets: {str(e)}")
      return {
          "success": False,
          "error": str(e),
          "project_id": "",
          "auth_provider": "azure",
      }

def safe_extract_response_content(response) -> str:
    """Safely extract content from LLM response without using eval()"""
    try:
        # Handle different response formats
        if hasattr(response, 'response'):
            return str(response.response)
        elif hasattr(response, 'chat_message') and hasattr(response.chat_message, 'content'):
            return str(response.chat_message.content)
        elif hasattr(response, 'content'):
            return str(response.content)
        elif isinstance(response, dict):
            if 'response' in response:
                return str(response['response'])
            elif 'content' in response:
                return str(response['content'])
        return str(response)
    except Exception as e:
        
        return str(response)

def generate_safe_uuid() -> str:
    """Generate a UUID string that won't cause octal parsing issues"""
    import uuid
    return str(uuid.uuid4())

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # You can restrict this to your frontend's origin
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health")
async def health_check():
    """Health check endpoint for Docker container monitoring"""
    return {"status": "healthy", "service": "agent-document-backend"}

# Configure logger
logger = logging.getLogger(__name__)

# ============================================================================
# REPOSITORY DOCUMENTATION ENDPOINTS
# ============================================================================

@app.get("/api/repos/list")
async def list_repositories():
    """
    List all available repositories from the external API
    
    Returns:
        RepoListResponse with list of repositories or error message
    """
    try:
        logger.info("Fetching repository list")
        response = await repo_service.list_repositories()
        
        # Return the response directly, including error cases
        # This allows frontend to display meaningful error messages
        return response
    except Exception as e:
        logger.error(f"Error listing repositories: {e}")
        return RepoListResponse(
            success=False,
            repos=[],
            error=f"Failed to list repositories: {str(e)}"
        )


@app.get("/api/repos/{repo_id}/chunks")
async def get_repo_chunks(repo_id: str):
    """
    Retrieve code chunks for a specific repository
    
    Args:
        repo_id: Repository ID to retrieve chunks for
        
    Returns:
        RepoChunksResponse with code chunks or error message
    """
    try:
        logger.info(f"Fetching chunks for repository: {repo_id}")
        response = await repo_service.retrieve_repo_chunks(repo_id)
        
        if not response.success:
            raise HTTPException(status_code=500, detail=response.error)
        
        return response
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error retrieving chunks for repo {repo_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to retrieve chunks: {str(e)}")


@app.post("/api/repos/generate-documentation")
async def generate_documentation(request: GenerateDocumentationRequest):
    """
    Generate comprehensive documentation for a repository
    
    Args:
        request: GenerateDocumentationRequest with repo_id and other parameters
        
    Returns:
        GenerateDocumentationResponse with generated documentation or error
    """
    try:
        logger.info(f"Generating documentation for repository: {request.repo_id}")
        
        # First, retrieve the repository chunks
        logger.info(f"Retrieving chunks for repository: {request.repo_id}")
        chunks_response = await repo_service.retrieve_repo_chunks(request.repo_id)
        
        if not chunks_response.success:
            raise HTTPException(status_code=500, detail=f"Failed to retrieve repo chunks: {chunks_response.error}")
        
        if not chunks_response.chunks:
            raise HTTPException(status_code=404, detail=f"No code chunks found for repository: {request.repo_id}")
        
        total_chunks = len(chunks_response.chunks)
        logger.info(f"Retrieved {total_chunks} chunks for repository: {request.repo_id}")
        
        # Provide helpful feedback for large repositories
        if total_chunks > 500:
            logger.info(f"Large repository detected ({total_chunks} chunks). Processing in batches for optimal performance.")
        
        # Generate documentation using the doc agent
        documentation = await doc_agent.generate_documentation(
            repo_id=request.repo_id,
            repo_name=request.repo_name,
            chunks=chunks_response.chunks,
            provider=request.llm_provider,
            model=request.llm_model
        )
        
        logger.info(f"Successfully generated documentation for repository: {request.repo_id}")
        
        return GenerateDocumentationResponse(
            success=True,
            repo_id=request.repo_id,
            repo_name=request.repo_name,
            documentation=documentation,
            error=None
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating documentation for repo {request.repo_id}: {e}")
        return GenerateDocumentationResponse(
            success=False,
            repo_id=request.repo_id,
            repo_name=request.repo_name,
            documentation="",
            error=f"Failed to generate documentation: {str(e)}"
        )


@app.post("/api/repos/chat")
async def chat_about_repo(request: dict):
    """
    Chat about a repository - ask questions and get answers based on the code
    
    Args:
        request: Dict with repo_id, message, repo_name, and optional conversation_history
        
    Returns:
        AgentResponse with the AI's answer
    """
    try:
        repo_id = request.get("repo_id")
        message = request.get("message")
        repo_name = request.get("repo_name")
        conversation_history = request.get("conversation_history", [])
        llm_provider = request.get("llm_provider")
        llm_model = request.get("llm_model")
        
        if not repo_id:
            raise HTTPException(status_code=400, detail="repo_id is required")
        if not message:
            raise HTTPException(status_code=400, detail="message is required")
        
        logger.info(f"Processing chat question for repository: {repo_id}")
        
        # Retrieve repository chunks
        chunks_response = await repo_service.retrieve_repo_chunks(repo_id)
        
        if not chunks_response.success:
            raise HTTPException(status_code=500, detail=f"Failed to retrieve repo chunks: {chunks_response.error}")
        
        if not chunks_response.chunks:
            raise HTTPException(status_code=404, detail=f"No code chunks found for repository: {repo_id}")
        
        logger.info(f"Chatting about repository: {llm_model} with message: {llm_provider}")
        
        # Generate response using the doc agent
        answer = await doc_agent.chat_about_repository(
            repo_id=repo_id,
            repo_name=repo_name,
            chunks=chunks_response.chunks,
            user_question=message,
            conversation_history=conversation_history,
            provider=llm_provider,
            model=llm_model
        )
        
        return AgentResponse(
            source="agent",
            content=answer
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing chat for repo: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process chat: {str(e)}")

# ============================================================================
# END REPOSITORY DOCUMENTATION ENDPOINTS
# ============================================================================
 


# General chat endpoint that classifies intent and routes to appropriate logic
@app.post("/execute-chat")
async def chat(request: ExecuteRequest):
    """
    General chat endpoint that classifies user intent and routes to appropriate test generation handlers.
    
    This is the main entry point for all chat interactions in the test automation system. It uses semantic 
    intent classification with keyword-based overrides to determine the user's intent and routes the request 
    to the appropriate test generation handler:
    
    - **informational/greeting/help**: Routes to general workflow execution for informational responses
    - **scenario**: Routes to test scenario generation handler for creating functional test scenarios
    - **scripts**: Routes to test script generation handler for creating automation test scripts
    - **case**: Routes to test case generation handler for creating detailed test cases
    - **epic/user story**: Routes to workflow execution or prompts user to specify generation type
    - **default**: Routes to general workflow execution for unclassified intents
    
    The endpoint supports conversation history context, session-based user story storage, and intelligent
    intent detection with fallback mechanisms. It prioritizes explicit keyword mentions for accurate routing.
    
    Args:
        request (ExecuteRequest): The chat request containing:
            - content (str): The user's message/content
            - message (str): Alternative message field (for backward compatibility)  
            - user_id (str): User identifier
            - tenant_id (str): Tenant identifier
            - session_id (str): Session identifier for maintaining context
            - llm_provider (str): LLM provider to use (e.g., "azure_openai", "aws_bedrock")
            - llm_model (str): Specific LLM model to use
            - userStoryText (str): Optional user story text content
            - conversationHistory (List[str]): Optional conversation history for context
            - test_case_option (str): Optional test case generation option
            - api_key (str): API key for LLM services
    
    Returns:
        JSONResponse: Response containing:
            - source (str): Response source ("agent" or "system")
            - content (str): The agent's response content (test scenarios, cases, scripts, or informational)
    
    Raises:
        HTTPException: If the request processing fails with appropriate error details including error type,
                      content, and source information
    
    Notes:
        - Uses SESSION_USER_STORIES dict to persist user stories across conversation turns
        - Supports conversation history analysis for better context understanding
        - Implements keyword-based intent overrides for accurate classification
        - Handles both single queries and conversation context for improved responses
    """
    try:
        # Handle both frontend and backend format
        query = request.content or request.message or ""
        llm_provider=request.llm_provider
        llm_model=request.llm_model
        intent = await classify_intent(query, llm_provider, llm_model)
        
        # Keyword-based overrides: check for cases and scenarios before scripts
        lower_query = (query or "").lower()
        # If the user explicitly mentions cases/scenario keywords prefer those intents
        if ('test case' in lower_query) or ('test cases' in lower_query) or (' case ' in lower_query) or lower_query.strip().endswith(' case') or lower_query.strip().startswith('case') or lower_query.strip() == 'case' or ' cases' in lower_query:
            intent = 'case'
        elif ('test scenario' in lower_query) or ('test scenarios' in lower_query) or (' scenario' in lower_query) or lower_query.strip().endswith(' scenario') or lower_query.strip().startswith('scenario') or lower_query.strip() == 'scenario':
            intent = 'scenario'
        # Fallback to scripts if the user mentions script-related words
        elif 'script' in lower_query or 'scripts' in lower_query or 'test script' in lower_query:
            intent = 'scripts'

        
        # Prepare conversation context for better agent responses
        conversation_context = ""
        user_story_from_history = ""
        
        # Extract user story from conversation history if available
        if request.conversationHistory and len(request.conversationHistory) > 0:
            # Look for user story content in conversation history
            for message in request.conversationHistory:
                if any(keyword in message.lower() for keyword in ['user story', 'user stories', 'heal-', 'acceptance criteria', 'feature name']):
                    user_story_from_history = message
                    break
            
            # Include recent conversation history for context
            recent_history = request.conversationHistory[-5:]  # Last 5 messages
            conversation_context = f"Previous conversation context:\n{chr(10).join(recent_history)}\n\nCurrent query: {query}"
        else:
            conversation_context = query
        
        # Use userStoryText from request, or from conversation history, or from session
        effective_user_story = request.userStoryText or user_story_from_history or SESSION_USER_STORIES.get(request.session_id or request.user_id, "")
        
        
        if intent == "informational" or intent == "greeting" or intent == "help":
            response = await execute_workflow(
                user_id=request.user_id,
                tenant_id=request.tenant_id,
                session_id=request.session_id,
                query=query,
                api_key=request.api_key,
                llm_provider=request.llm_provider,
                llm_model=request.llm_model,
                intent="informational"
            )
            
            return JSONResponse(content=[{"source": "agent", "content": response["response"]}], status_code=200, media_type="application/json")
            # return JSONResponse(content=[{"source": response.source, "content": response.content}], status_code=200, media_type="application/json")
       
        else:
            response = await execute_workflow(
                user_id=request.user_id,
                tenant_id=request.tenant_id,
                session_id=request.session_id,
                query=conversation_context,
                api_key=request.api_key,
                llm_provider=request.llm_provider,
                llm_model=request.llm_model,
            )
            
            return JSONResponse(content=[{"source": response.source, "content": response.content}], status_code=200, media_type="application/json")
    except Exception as e:
        error_message = {
            "type": "error",
            "content": f"Error: {str(e)}",
            "source": "system"
        }
        raise HTTPException(status_code=500, detail=error_message) from e


# Allowed text file extensions for extraction from zip
TEXT_FILE_EXTENSIONS = {
    '.txt', '.md', '.py', '.js', '.ts', '.tsx', '.jsx', '.json', '.yaml', '.yml',
    '.xml', '.html', '.css', '.sql', '.sh', '.bat', '.ps1', '.c', '.cpp', '.h',
    '.java', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala', '.r', '.m',
    '.csv', '.ini', '.cfg', '.conf', '.env', '.gitignore', '.dockerignore',
    '.editorconfig', '.eslintrc', '.prettierrc', '.babelrc'
}


def extract_text_from_zip(zip_bytes: bytes, max_file_size: int = 1024 * 1024) -> str:
    """
    Extract text content from all text files within a zip archive.
    
    Args:
        zip_bytes: The zip file content as bytes
        max_file_size: Maximum size per file to extract (default 1MB)
    
    Returns:
        Concatenated text content from all text files in the zip
    """
    extracted_content = []
    
    with zipfile.ZipFile(io.BytesIO(zip_bytes), 'r') as zf:
        for file_info in zf.infolist():
            # Skip directories
            if file_info.is_dir():
                continue
            
            # Get file extension
            file_name = file_info.filename
            _, ext = os.path.splitext(file_name.lower())
            
            # Skip non-text files
            if ext not in TEXT_FILE_EXTENSIONS:
                continue
            
            # Skip large files
            if file_info.file_size > max_file_size:
                extracted_content.append(f"--- File: {file_name} (skipped: too large, {file_info.file_size} bytes) ---")
                continue
            
            try:
                with zf.open(file_info) as f:
                    content = f.read()
                    # Try to decode as UTF-8, fall back to latin-1
                    try:
                        text_content = content.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            text_content = content.decode('latin-1')
                        except UnicodeDecodeError:
                            extracted_content.append(f"--- File: {file_name} (skipped: encoding error) ---")
                            continue
                    
                    extracted_content.append(f"--- File: {file_name} ---\n{text_content}\n--- End of {file_name} ---")
            except Exception as e:
                extracted_content.append(f"--- File: {file_name} (error: {str(e)}) ---")
    
    return "\n\n".join(extracted_content)


@app.post("/execute-chat-with-zip")
async def chat_with_zip(
    message: str = Form(""),
    llm_provider: str = Form("aws_bedrock"),
    llm_model: str = Form("anthropic.claude-3-5-sonnet-20241022-v2:0"),
    user_id: str = Form("anonymous"),
    tenant_id: str = Form("default"),
    session_id: str = Form(""),
    zip_file: UploadFile = File(None),
):
    """
    Chat endpoint that accepts a zip file upload, extracts text content from files,
    and sends to LLM for processing.
    
    Args:
        message: The user's message/query
        llm_provider: LLM provider to use
        llm_model: LLM model to use
        user_id: User identifier
        tenant_id: Tenant identifier
        session_id: Session identifier
        zip_file: The uploaded zip file
    
    Returns:
        JSONResponse with agent response
    """
    try:
        extracted_content = ""
        
        # Extract content from zip file if provided
        if zip_file and zip_file.filename:
            if not zip_file.filename.lower().endswith('.zip'):
                raise HTTPException(status_code=400, detail="Only .zip files are allowed")
            
            zip_bytes = await zip_file.read()
            extracted_content = extract_text_from_zip(zip_bytes)
            
            if not extracted_content.strip():
                extracted_content = "(No text files found in the zip archive)"
        
        # Build the full query with extracted content
        if extracted_content:
            full_query = f"{message}\n\n--- Uploaded Zip Contents ---\n{extracted_content}"
        else:
            full_query = message
        
        if not full_query.strip():
            raise HTTPException(status_code=400, detail="No message or file content provided")
        
        # Classify intent and process
        intent = await classify_intent(full_query, llm_provider, llm_model)
        
        # Route to appropriate handler
        if intent in ["informational", "greeting", "help"]:
            response = await execute_workflow(
                user_id=user_id,
                tenant_id=tenant_id,
                session_id=session_id,
                query=full_query,
                api_key="",
                llm_provider=llm_provider,
                llm_model=llm_model,
                intent="informational"
            )
            return JSONResponse(
                content=[{"source": "agent", "content": response.get("response", str(response))}],
                status_code=200,
                media_type="application/json"
            )
        else:
            response = await execute_workflow(
                user_id=user_id,
                tenant_id=tenant_id,
                session_id=session_id,
                query=full_query,
                api_key="",
                llm_provider=llm_provider,
                llm_model=llm_model,
            )
            return JSONResponse(
                content=[{"source": "agent", "content": response.get("response", str(response))}],
                status_code=200,
                media_type="application/json"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in chat_with_zip: {str(e)}")
        error_message = {
            "type": "error",
            "content": f"Error processing request: {str(e)}",
            "source": "system"
        }
        raise HTTPException(status_code=500, detail=error_message) from e

 
# SharePoint file listing endpoints
@app.get("/sharepoint-files", response_model=SharePointListResponse)
async def list_sharepoint_files(
    folder_path: str = "") -> SharePointListResponse:
    """List files from SharePoint"""
    try:
        # SharePointConfig will use secrets_manager global context
        config = SharePointConfig()
        
        # Validate SharePoint configuration
        if not config.sharepoint_folder:
            raise ValueError("SharePoint is not configured for this project.")
        
        client = SharePointClient(
            client_id=config.client_id,
            client_secret=config.client_secret,
            tenant_id=config.tenant_id,
            site_url=config.site_url,
            sharepoint_folder=config.sharepoint_folder,
            logger=logger
        )
        logger.info(f"Listing SharePoint files for site: {config.sharepoint_folder}")
        # Get the sharepoint site-folder from configuration
        sharepoint_site_folder = config.sharepoint_folder
        
        # Parse the site name and folder path from the configuration
        if "/" in sharepoint_site_folder:
            env_site_name, env_folder_path = sharepoint_site_folder.split("/", 1)
        else:
            env_site_name = sharepoint_site_folder
            env_folder_path = ""
        
        # Use configuration values for site, but allow folder_path parameter for navigation
        actual_site_name = env_site_name
        # If folder_path is provided in query params, use it; otherwise use configured default
        actual_folder_path = folder_path if folder_path else env_folder_path

        # Get site drives
        drives_response = client.get_site_drives(actual_site_name)
        if not drives_response.get('value'):
            raise Exception("No drives found for the site")
        
        # Use the first available drive (usually Documents)
        drive_id = drives_response['value'][0]['id']
        
        # Get folder contents
        files = client.get_folder_contents(drive_id, actual_folder_path)
        
        # Format files for frontend
        formatted_files = []
        for file_item in files:
            formatted_file = SharePointFileItem(
                id=file_item.get('id', ''),
                name=file_item.get('name', ''),
                isFolder=bool(file_item.get('folder')),
                size=file_item.get('size', 0),
                createdDateTime=file_item.get('createdDateTime', ''),
                lastModifiedDateTime=file_item.get('lastModifiedDateTime', ''),
                webUrl=file_item.get('webUrl', ''),
                downloadUrl=file_item.get('@microsoft.graph.downloadUrl', ''),
                mimeType=file_item.get('file', {}).get('mimeType', '') if file_item.get('file') else 'folder',
                createdBy=file_item.get('createdBy', {}).get('user', {}).get('displayName', 'Unknown')
            )
            formatted_files.append(formatted_file)
        
        return SharePointListResponse(
            success=True,
            files=formatted_files,
            folder_path=actual_folder_path,
            site_name=actual_site_name
        )
        
    except ValueError as e:
        logger.error(f"SharePoint configuration error: {str(e)}")
        return SharePointListResponse(
            success=False,
            error="SharePoint is not configured for this project.",
            files=[]
        )
    except Exception as e:
        logger.error(f"Error listing SharePoint files: {str(e)}")
        return SharePointListResponse(
            success=False,
            error=f"Failed to list SharePoint files: {str(e)}",
            files=[]
        )

@app.get("/sharepoint-files/{file_id}/download")
async def download_sharepoint_file(
    file_id: str,
    project_id: str = "",
    auth_provider: str = "",
    site_name: str = "",
    folder_path: str = ""
):
    """Download a specific file from SharePoint
     This endpoint is restricted to direct browser access and manual invocation only.
    Not available for MCP workflow execution."""
    try:
        # SharePointConfig will use secrets_manager global context
        config = SharePointConfig()
        client = SharePointClient(
            client_id=config.client_id,
            client_secret=config.client_secret,
            tenant_id=config.tenant_id,
            site_url=config.site_url,
            sharepoint_folder=config.sharepoint_folder,
            logger=logger
        )
        
        # Get the sharepoint site-folder from environment variable
        sharepoint_site_folder = config.sharepoint_folder
        
        # Parse the site name and folder path from the environment variable
        if "/" in sharepoint_site_folder:
            env_site_name, env_folder_path = sharepoint_site_folder.split("/", 1)
        else:
            env_site_name = sharepoint_site_folder
            env_folder_path = ""
        
        # Use environment values
        actual_site_name = env_site_name
        actual_folder_path = env_folder_path

        # Get site drives
        drives_response = client.get_site_drives(actual_site_name)
        if not drives_response.get('value'):
            raise Exception("No drives found for the site")
        
        # Use the first available drive (usually Documents)
        drive_id = drives_response['value'][0]['id']
        
        # Get file metadata
        file_metadata = client.get_file_metadata(drive_id, file_id)
        
        # Get file content using download URL
        import requests
        download_url = file_metadata.get('@microsoft.graph.downloadUrl')
        if not download_url:
            raise Exception("Download URL not available for this file")
        
        response = requests.get(download_url)
        response.raise_for_status()
        
        from fastapi.responses import Response
        
        # Return file content with appropriate headers
        return Response(
            content=response.content,
            media_type=file_metadata.get('file', {}).get('mimeType', 'application/octet-stream'),
            headers={
                "Content-Disposition": f"attachment; filename=\"{file_metadata.get('name', 'download')}\""
            }
        )
        
    except Exception as e:
        logger.error(f"Error downloading SharePoint file {file_id}: {str(e)}")
        return JSONResponse(
            content={
                "success": False,
                "error": f"Failed to download SharePoint file: {str(e)}"
            },
            status_code=500
        )

# Download method for MCP server compatibility - keep as POST for body support
@app.post("/sharepoint/download", response_model=SharePointDownloadResponse)
async def download_sharepoint_file_mcp(
    requestModel : SharePointDownloadRequest
) -> SharePointDownloadResponse:
    """Download a specific file from SharePoint - Works with MCP server"""
    try:
        # SharePointConfig will use secrets_manager global context
        config = SharePointConfig()
        client = SharePointClient(
            client_id=config.client_id,
            client_secret=config.client_secret,
            tenant_id=config.tenant_id,
            site_url=config.site_url,
            sharepoint_folder=config.sharepoint_folder,
            logger=logger
        )
        # Get the sharepoint site-folder from config
        sharepoint_site_folder = config.sharepoint_folder
        
        # Parse the site name and folder path from the environment variable
        if "/" in sharepoint_site_folder:
            env_site_name, env_folder_path = sharepoint_site_folder.split("/", 1)
        else:
            env_site_name = sharepoint_site_folder
            env_folder_path = "BuilderAI"
        
        # Use environment values
        actual_site_name = env_site_name
        actual_folder_path = env_folder_path

        # Get site drives
        drives_response = client.get_site_drives(actual_site_name)
        if not drives_response.get('value'):
            raise Exception("No drives found for the site")
        
        # Use the first available drive (usually Documents)
        drive_id = drives_response['value'][0]['id']
        
        # Get file metadata
        file_metadata = client.get_file_metadata(drive_id, requestModel.file_id)
        
        # Download file content using download URL from metadata
        import requests
        import base64
        
        download_url = file_metadata.get('@microsoft.graph.downloadUrl')
        if not download_url:
            raise Exception("Download URL not available for this file")
        
        response = requests.get(download_url)
        response.raise_for_status()
        file_content = response.content
        
        # Encode binary content as base64 for JSON serialization
        file_content_b64 = base64.b64encode(file_content).decode('utf-8')
        
        return SharePointDownloadResponse(
            success=True,
            file_metadata=SharePointFileMetadata(
                id=file_metadata.get('id'),
                name=file_metadata.get('name', 'download'),
                size=file_metadata.get('size', 0),
                mime_type=file_metadata.get('file', {}).get('mimeType', 'application/octet-stream'),
                created_date=file_metadata.get('createdDateTime'),
                modified_date=file_metadata.get('lastModifiedDateTime'),
                web_url=file_metadata.get('webUrl')
            ),
            file_content_base64=file_content_b64,
            content_length=len(file_content)
        )
    except Exception as e:
        logger.error(f"Error downloading SharePoint file {requestModel.file_id}: {str(e)}")
        return SharePointDownloadResponse(
            success=False,
            error=f"Failed to download SharePoint file: {str(e)}"
        )

# SharePoint upload endpoint - handles both text content and binary files
@app.post("/upload-to-sharepoint")
async def upload_to_sharepoint(
    file_content: str = Form(None),
    file_name: str = Form(...),
    folder_path: str = Form(default=""),
    site_name: str = Form(default=""),
    file: UploadFile = File(None),
    project_id: str = Form(""),
    auth_provider: str = Form(""),
):
    """ This endpoint is restricted to direct browser access and manual invocation only.
    Not available for MCP workflow execution."""
    try:
        # SharePointConfig will use secrets_manager global context
        config = SharePointConfig()
        client = SharePointClient(
            client_id=config.client_id,
            client_secret=config.client_secret,
            tenant_id=config.tenant_id,
            site_url=config.site_url,
            sharepoint_folder=config.sharepoint_folder,
            logger=logger
        )
         # Get the sharepoint site-folder from config
        sharepoint_site_folder = config.sharepoint_folder
        
        # Parse the site name and folder path from the environment variable
        if "/" in sharepoint_site_folder:
            env_site_name, env_folder_path = sharepoint_site_folder.split("/", 1)
        else:
            env_site_name = sharepoint_site_folder
            env_folder_path = ""
        
        # Use provided values or fall back to environment variable values
        actual_site_name =  env_site_name
        actual_folder_path = env_folder_path

        # Get site drives
        drives_response = client.get_site_drives(actual_site_name)
        if not drives_response.get('value'):
            raise Exception("No drives found for the site")
        
        # Use the first available drive (usually Documents)
        drive_id = drives_response['value'][0]['id']
        
        # Handle file content - either from uploaded file or text content
        if file:
            # Handle uploaded file (Excel, DOCX, etc.)
            file_bytes = await file.read()
        elif file_content:
            # Handle text content - convert to DOCX
            file_bytes = text_to_docx_bytes(file_content)
        else:
            raise Exception("No file or content provided")
            
        # Upload file to SharePoint
        upload_result = client.upload_file(
            drive_id=drive_id,
            folder_path=actual_folder_path,
            file_name=file_name,
            file_content=file_bytes
        )
        
        # Construct SharePoint URLs for file access
        base_site_url = f"https://wipro365.sharepoint.com/sites/{actual_site_name}"
        
        # Properly encode folder path and file name for URLs
        import urllib.parse
        encoded_folder_path = urllib.parse.quote(actual_folder_path, safe='')
        encoded_file_name = urllib.parse.quote(file_name, safe='')
        
        # Construct the complete file path for SharePoint
        file_server_relative_url = f"/sites/{actual_site_name}/Shared Documents/{actual_folder_path}/{file_name}"
        encoded_file_path = urllib.parse.quote(file_server_relative_url, safe='')
        
        # SharePoint folder view URL - opens the folder containing the file
        folder_view_url = f"{base_site_url}/Shared%20Documents/Forms/AllItems.aspx?id=%2Fsites%2F{actual_site_name}%2FShared%20Documents%2F{encoded_folder_path}&viewid=00000000-0000-0000-0000-000000000000"
        
        # SharePoint file view URL - opens the specific file in SharePoint's web viewer
        content_view_url = f"{base_site_url}/Shared%20Documents/Forms/AllItems.aspx?id=%2Fsites%2F{actual_site_name}%2FShared%20Documents%2F{encoded_folder_path}"

        
        return JSONResponse(
            content={
                "success": True,
                "message": f"File '{file_name}' uploaded successfully to SharePoint",
                "file_url": content_view_url,  # Primary URL to open file in SharePoint viewer
                "folder_url": folder_view_url,  # URL to open the folder containing the file
                "download_url": upload_result.get("webUrl", ""),  # Keep original download URL as backup
                "file_id": upload_result.get("id", ""),
            },
            status_code=200
        )
        
    except Exception as e:
        logger.error(f"Error uploading to SharePoint: {str(e)}")
        return JSONResponse(
            content={
                "success": False,
                "error": f"Failed to upload to SharePoint: {str(e)}"
            },
            status_code=500
        )
    

# MCP compatible upload to sharepoint function
@app.post("/sharepoint/upload", response_model=SharePointUploadResponse)
async def upload_to_sharepoint_mcp(request: SharePointUploadRequest) -> SharePointUploadResponse:
    """Upload document to SharePoint - MCP compatible version"""
    try:
        # SharePointConfig will use secrets_manager global context
        config = SharePointConfig()
        client = SharePointClient(
            client_id=config.client_id,
            client_secret=config.client_secret,
            tenant_id=config.tenant_id,
            site_url=config.site_url,
            sharepoint_folder=config.sharepoint_folder,
            logger=logger
        )
        
        # Get the sharepoint site-folder from config
        sharepoint_site_folder = config.sharepoint_folder
        
        # Parse the site name and folder path from the environment variable
        if "/" in sharepoint_site_folder:
            env_site_name, env_folder_path = sharepoint_site_folder.split("/", 1)
        else:
            env_site_name = sharepoint_site_folder
            env_folder_path = ""
        
        # Use request model values or fallback to environment values
        actual_site_name = request.site_name or env_site_name
        actual_folder_path = request.folder_path or env_folder_path

        # Get site drives
        drives_response = client.get_site_drives(actual_site_name)
        if not drives_response.get('value'):
            raise Exception("No drives found for the site")
        
        drive_id = drives_response['value'][0]['id']
        
        # Use request model attributes
        file_name = request.file_name
        file_content = request.file_content
        
        # Convert text content to DOCX bytes for SharePoint upload
        file_bytes = text_to_docx_bytes(file_content)
            
        # Upload file to SharePoint
        upload_result = client.upload_file(
            drive_id=drive_id,
            folder_path=actual_folder_path,
            file_name=file_name,
            file_content=file_bytes
        )
        
        # Construct SharePoint URLs for file access
        base_site_url = f"https://wipro365.sharepoint.com/sites/{actual_site_name}"
        
        # Properly encode folder path and file name for URLs
        import urllib.parse
        encoded_folder_path = urllib.parse.quote(actual_folder_path, safe='')
        encoded_file_name = urllib.parse.quote(file_name, safe='')
        
        # Construct the complete file path for SharePoint
        file_server_relative_url = f"/sites/{actual_site_name}/Shared Documents/{actual_folder_path}/{file_name}"
        encoded_file_path = urllib.parse.quote(file_server_relative_url, safe='')
        
        # SharePoint folder view URL - opens the folder containing the file
        folder_view_url = f"{base_site_url}/Shared%20Documents/Forms/AllItems.aspx?id=%2Fsites%2F{actual_site_name}%2FShared%20Documents%2F{encoded_folder_path}&viewid=00000000-0000-0000-0000-000000000000"
        
        # SharePoint file view URL - opens the specific file in SharePoint's web viewer
        content_view_url = f"{base_site_url}/Shared%20Documents/Forms/AllItems.aspx?id=%2Fsites%2F{actual_site_name}%2FShared%20Documents%2F{encoded_folder_path}"
        
        # Return MCP-compatible response using the schema model
        return SharePointUploadResponse(
            success=True,
            message=f"File '{file_name}' uploaded successfully to SharePoint",
            file_id=upload_result.get("id", ""),
            web_url=upload_result.get("webUrl", ""),
            file_url=content_view_url,  # Primary URL to open file in SharePoint viewer
            folder_url=folder_view_url,  # URL to open the folder containing the file
            site_name=actual_site_name,
            folder_path=actual_folder_path
        )
        
    except Exception as e:
        logger.error(f"Error uploading to SharePoint: {str(e)}")
        return SharePointUploadResponse(
            success=False,
            message="Upload failed",
            error=f"Failed to upload to SharePoint: {str(e)}"
        )
      
def text_to_docx_bytes(text: str) -> bytes:
    """Convert text to DOCX bytes with proper formatting. Falls back to plain text if docx module not available."""
    if Document is None:
        # Fallback to plain text
        return text.encode('utf-8')
    
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Create paragraphs for each line with better formatting
    paragraphs = []
    
    # Add AI generated disclaimer at the top in red small fonts
    ai_disclaimer = doc.add_paragraph()
    ai_disclaimer_run = ai_disclaimer.add_run('This is AI generated content')
    ai_disclaimer_run.font.size = Pt(9)  # Small font
    ai_disclaimer_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
    ai_disclaimer_run.italic = True
    ai_disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ai_disclaimer.space_after = Pt(6)
    
    # Add title - centered and bold
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run('Documentation')
    title_run.bold = True
    title_run.font.size = Pt(16)  # 16pt font size
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space after title
    title_paragraph.space_after = Pt(12)
    
    # Add content
    for line in text.split('\n'):
        if line.strip():  # Only add non-empty lines
            paragraph = doc.add_paragraph(line)
            paragraph.space_after = Pt(6)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()

# Agent prompts - FastMCP-friendly endpoint
@app.post("/get-agent-prompts", response_model=AgentPromptsResponse)
async def get_agent_prompts(request: AgentPromptsRequest):
    """
    Retrieve prompts for a given agent from the external Agent service.

    This FastMCP-friendly endpoint accepts an `agent_id` and forwards the request
    to the external service at:
    {BUILDAI_BACKEND_API_URL}/v1.0/agents/{agent_id}/prompts

    SSL verification can be controlled via the VERIFY_SSL secret/environment variable
    (set to 'true' to enable verification).
    """
    try:
        agent_id = request.agent_id
        if not agent_id:
            return AgentPromptsResponse(
                success=False, prompts=[], count=0, error="agent_id is required"
            )

        external_api_url = (
            f"{BUILDAI_BACKEND_API_URL}/v1.0/agents/{agent_id}/prompts"
        )

        verify_ssl = get_secret("VERIFY_SSL", None, "false").lower() == "true"
        if not verify_ssl:
            logger.warning(
                "SSL verification is disabled for external Agent API requests. Enable VERIFY_SSL=true for production."
            )

        async with httpx.AsyncClient(timeout=30.0, verify=verify_ssl) as client:
            resp = await client.get(external_api_url)
            resp.raise_for_status()
            external_data = resp.json()

        # If external response matches expected format, return as-is (validated by model)
        if (
            isinstance(external_data, dict)
            and external_data.get("status") is not None
            and isinstance(external_data.get("data"), dict)
        ):
            # Ensure prompts are list[str]
            data_obj = external_data["data"]
            prompts_val = data_obj.get("prompts") or []
            if not isinstance(prompts_val, list):
                prompts_val = [str(prompts_val)]

            return AgentPromptsResponse(
                status=str(external_data.get("status")),
                message=str(external_data.get("message", "")),
                data={
                    "agent_id": data_obj.get("agent_id", request.agent_id),
                    "agent_name": data_obj.get("agent_name"),
                    "prompts": [str(p) for p in prompts_val],
                    "total_prompts": int(data_obj.get("total_prompts", len(prompts_val)))
                }
            )

        # Fallback normalization if structure differs
        if isinstance(external_data, list):
            prompts = [str(p) for p in external_data]
        elif isinstance(external_data, dict) and "prompts" in external_data:
            raw_prompts = external_data.get("prompts") or []
            prompts = [str(p) for p in raw_prompts] if isinstance(raw_prompts, list) else [str(raw_prompts)]
        else:
            prompts = []

        return AgentPromptsResponse(
            status="success",
            message=f"Retrieved {len(prompts)} prompts for agent successfully",
            data={
                "agent_id": request.agent_id,
                "agent_name": None,
                "prompts": prompts,
                "total_prompts": len(prompts)
            }
        )

    except httpx.TimeoutException:
        logger.error("Timeout when calling external Agent API")
        return AgentPromptsResponse(status="error", message="Request timeout when fetching agent prompts", data={"agent_id": request.agent_id, "agent_name": None, "prompts": [], "total_prompts": 0})
    except httpx.HTTPStatusError as e:
        logger.error(f"HTTP error when calling external Agent API: {e.response.status_code}")
        return AgentPromptsResponse(status="error", message=f"External API error: {e.response.status_code}", data={"agent_id": request.agent_id, "agent_name": None, "prompts": [], "total_prompts": 0})
    except httpx.ConnectError as e:
        logger.error(f"Connection error when calling external Agent API: {str(e)}")
        return AgentPromptsResponse(status="error", message="Unable to connect to external Agent API", data={"agent_id": request.agent_id, "agent_name": None, "prompts": [], "total_prompts": 0})
    except Exception as e:
        error_msg = str(e)
        if "certificate verify failed" in error_msg.lower() or "ssl" in error_msg.lower():
            logger.error(f"SSL certificate error when calling external Agent API: {error_msg}")
            return AgentPromptsResponse(status="error", message="SSL certificate verification failed", data={"agent_id": request.agent_id, "agent_name": None, "prompts": [], "total_prompts": 0})
        logger.error(f"Error fetching agent prompts: {error_msg}")
        return AgentPromptsResponse(status="error", message=f"Failed to fetch agent prompts: {error_msg}", data={"agent_id": request.agent_id, "agent_name": None, "prompts": [], "total_prompts": 0})


# ============================================================================
# CONFLUENCE UPLOAD ENDPOINT
# ============================================================================
# COMMENTED OUT: Confluence output not needed

# @app.post("/upload-to-confluence", response_model=UploadToConfluenceResponse)
# async def upload_to_confluence(request: UploadToConfluenceRequest):
#     """Upload/create a new page in Confluence or update if it already exists"""
#     logger.info(f"🔍 ENDPOINT HIT: /upload-to-confluence - Title: {request.page_title}")
#     try:
#         # Initialize Confluence client using global secrets context
#         config = ConfluenceConfig()
#         client = ConfluenceClient(
#             base_url=config.base_url,
#             email=config.email,
#             api_key=config.api_key,
#             space_key=config.space_key,
#             logger=logger
#         )
#         
#         # Use provided space key or fall back to configured one
#         target_space_key = request.space_key or config.space_key
#         
#         logger.info(f"Creating/updating page '{request.page_title}' in space '{target_space_key}'")
#         
#         # Create or update the page (if it already exists, it will be updated)
#         result_page = client.create_or_update_page(
#             title=request.page_title,
#             content=request.content,
#             parent_id=request.parent_id,
#             space_key=target_space_key
#         )
#         
#         # Extract page information
#         page_id = result_page.get('id', '')
#         space_info = result_page.get('space', {})
#         
#         # Construct web URL
#         web_url = f"{config.base_url}/wiki{result_page.get('_links', {}).get('webui', '')}"
#         
#         logger.info(f"Successfully created/updated Confluence page: {request.page_title} (ID: {page_id})")
#         
#         return UploadToConfluenceResponse(
#             success=True,
#             page_id=page_id,
#             page_title=result_page.get('title', request.page_title),
#             space_key=space_info.get('key', target_space_key),
#             space_name=space_info.get('name', ''),
#             web_url=web_url,
#             created_date=result_page.get('createdDate', ''),
#             message=f"Page '{request.page_title}' has been successfully created/updated in Confluence."
#         )
#         
#     except ValueError as e:
#         logger.error(f"Confluence configuration error: {str(e)}")
#         return UploadToConfluenceResponse(
#             success=False,
#             error="Confluence is not configured for this project."
#         )
#     except Exception as e:
#         logger.error(f"Error creating Confluence page: {str(e)}")
#         return UploadToConfluenceResponse(
#             success=False,
#             error=f"Failed to create Confluence page: {str(e)}"
#         )


# ============================================================================
# CONFLUENCE LISTING ENDPOINTS
# ============================================================================

# ============================================================================
# CONFLUENCE LISTING ENDPOINTS
# ============================================================================
# COMMENTED OUT: Confluence functionality not needed

# @app.get("/confluence-pages", response_model=ConfluencePagesResponse)
# async def list_confluence_pages(space_key: Optional[str] = None):
#     """List Confluence pages from spaces"""
#     logger.info("🔍 ENDPOINT HIT: /confluence-pages")
#     try:
#         # Initialize Confluence client using global secrets context
#         config = ConfluenceConfig()
#         client = ConfluenceClient(
#             base_url=config.base_url,
#             email=config.email,
#             api_key=config.api_key,
#             space_key=config.space_key,
#             logger=logger
#         )
#         
#         # Get space key from environment variable or use provided value
#         env_space_key = config.space_key
#         target_space_key = space_key if space_key else env_space_key
#         logger.info(f"Fetching all pages (including child pages) from space: {target_space_key}")
#         
#         # Fetch all pages from Confluence (with pagination to get all pages including children)
#         confluence_pages = client.list_pages(space_key=target_space_key, limit=100, fetch_all=True)
#         
#         logger.info(f"Total pages fetched: {len(confluence_pages)}")
#         
#         # Format pages for response
#         formatted_pages = []
#         for page in confluence_pages:
#             space_info = page.get('space', {})
#             version_info = page.get('version', {})
#             ancestors = page.get('ancestors', [])
#             
#             # Get parent page info if exists
#             parent_title = ancestors[-1].get('title', '') if ancestors else ''
#             parent_id = ancestors[-1].get('id', '') if ancestors else ''
#             
#             formatted_page = {
#                 "id": page.get('id', ''),
#                 "title": page.get('title', 'Untitled'),
#                 "spaceKey": space_info.get('key', ''),
#                 "spaceName": space_info.get('name', ''),
#                 "createdDate": page.get('createdDate', ''),
#                 "lastModifiedDate": version_info.get('when', page.get('createdDate', '')),
#                 "webUrl": f"{config.base_url}/wiki{page.get('_links', {}).get('webui', '')}",
#                 "contentUrl": page.get('_links', {}).get('self', ''),
#                 "parentTitle": parent_title,
#                 "parentId": parent_id,
#                 "depth": len(ancestors)  # 0 = root page, 1+ = child pages
#             }
#             formatted_pages.append(formatted_page)
#         
#         return ConfluencePagesResponse(
#             success=True,
#             pages=formatted_pages
#         )
#         
#     except ValueError as e:
#         logger.error(f"Confluence configuration error: {str(e)}")
#         return ConfluencePagesResponse(
#             success=False,
#             error="Confluence is not configured for this project.",
#             pages=[]
#         )
#     except Exception as e:
#         logger.error(f"Error listing Confluence pages: {str(e)}")
#         return ConfluencePagesResponse(
#             success=False,
#             error=f"Failed to list Confluence pages: {str(e)}",
#             pages=[]
#         )


# @app.get("/confluence-page-content/{page_id}", response_model=ConfluencePageContentResponse)
# async def get_confluence_page_content(page_id: str = ""):
#     """Get content from a Confluence page"""
#     logger.info(f"🔍 ENDPOINT HIT: /confluence-page-content/{page_id}")
#     try:
#         # Initialize Confluence client using global secrets context
#         config = ConfluenceConfig()
#         client = ConfluenceClient(
#             base_url=config.base_url,
#             email=config.email,
#             api_key=config.api_key,
#             space_key=config.space_key,
#             logger=logger
#         )
#         
#         # Fetch page content
#         page_data = client.get_page_content(page_id)
#         
#         # Extract text content from storage format
#         body_storage = page_data.get('body', {}).get('storage', {})
#         storage_content = body_storage.get('value', '')
#         extracted_content = client.extract_text_from_storage(storage_content)
#         
#         # Get page metadata
#         space_info = page_data.get('space', {})
#         version_info = page_data.get('version', {})
#         
#         return ConfluencePageContentResponse(
#             success=True,
#             page_id=page_id,
#             page_title=page_data.get('title', 'Untitled'),
#             content=extracted_content,
#             space_key=space_info.get('key', ''),
#             space_name=space_info.get('name', ''),
#             last_modified_date=version_info.get('when', page_data.get('createdDate', '')),
#             web_url=f"{config.base_url}/wiki{page_data.get('_links', {}).get('webui', '')}"
#         )
#         
#     except ValueError as e:
#         logger.error(f"Confluence configuration error: {str(e)}")
#         return ConfluencePageContentResponse(
#             success=False,
#             error="Confluence is not configured for this project."
#         )
#     except Exception as e:
#         logger.error(f"Error extracting Confluence page content {page_id}: {str(e)}")
#         return ConfluencePageContentResponse(
#             success=False,
#             error=f"Failed to extract Confluence page content: {str(e)}"
#         )


# Example usage
if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8002)
