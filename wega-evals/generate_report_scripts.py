"""Generate a professional DOCX evaluation report for the testcase_to_scripts agent."""

from __future__ import annotations

import io
import os
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# ── Evaluation Data ──────────────────────────────────────────────────

AGENT_NAME = "testcase_to_scripts"
AGENT_DESCRIPTION = (
    "Converts structured test cases into executable test automation scripts. "
    "The agent supports multiple frameworks (Playwright, Selenium BDD) and "
    "languages (JavaScript, Java), generating complete test scripts with proper "
    "assertions, data handling, and framework-specific patterns."
)
EVAL_DATE = "May 16, 2026"
DATASET = "tc-to-scripts-eval"
ITEMS_EVALUATED = 100
PASS_THRESHOLD = 0.70
JUDGE_MODEL = "gemini-2.5-flash (Google Vertex AI)"

# Combined results
DIMENSIONS = {
    # Domain-specific
    "script_correctness":      {"score": 0.57, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    "framework_compliance":    {"score": 0.69, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    "test_coverage_mapping":   {"score": 0.48, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    "code_quality":            {"score": 0.66, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    "executability":           {"score": 0.54, "category": "Domain Quality",  "evaluator": "LLM Judge"},
    # Common LLM Judge
    "coherence":               {"score": 0.73, "category": "Output Quality",  "evaluator": "LLM Judge"},
    "conciseness":             {"score": 0.37, "category": "Output Quality",  "evaluator": "LLM Judge"},
    "consistency":             {"score": 0.62, "category": "Output Quality",  "evaluator": "LLM Judge"},
    "faithfulness":            {"score": 0.32, "category": "Output Quality",  "evaluator": "LLM Judge"},
    # Safety
    "hallucination":           {"score": 0.18, "category": "Safety & Trust",  "evaluator": "LLM Judge"},
    "toxicity":                {"score": 1.00, "category": "Safety & Trust",  "evaluator": "LLM Judge"},
    "data_privacy_compliance": {"score": 0.66, "category": "Safety & Trust",  "evaluator": "LLM Judge"},
    # Compliance
    "policy_compliance":       {"score": 0.54, "category": "Compliance",      "evaluator": "LLM Judge"},
    # Performance & Programmatic
    "structural_validity":     {"score": 1.00, "category": "Structural",      "evaluator": "LLM Judge"},
    "latency":                 {"score": 0.35, "category": "Performance",     "evaluator": "Programmatic"},
    "cost_efficiency":         {"score": 0.67, "category": "Performance",     "evaluator": "Programmatic"},
    "token_efficiency":        {"score": 0.57, "category": "Performance",     "evaluator": "Programmatic"},
}

# ── Colour palette ───────────────────────────────────────────────────

BRAND_DARK  = RGBColor(0x1B, 0x2A, 0x4A)   # dark navy
BRAND_MID   = RGBColor(0x2E, 0x86, 0xAB)   # teal
BRAND_LIGHT = RGBColor(0xE8, 0xF1, 0xF5)   # light blue-grey
PASS_GREEN  = RGBColor(0x27, 0xAE, 0x60)
FAIL_RED    = RGBColor(0xE7, 0x4C, 0x3C)
GREY        = RGBColor(0x7F, 0x8C, 0x8D)


def _set_cell_shading(cell, color_hex: str):
    """Set background shading for a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def _add_styled_heading(doc, text, level=1):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = BRAND_DARK
    return heading


def _create_bar_chart(dimensions: dict, filename: str):
    """Create a horizontal bar chart of dimension scores."""
    sorted_dims = sorted(dimensions.items(), key=lambda x: x[1]["score"], reverse=True)
    names = [d[0].replace("_", " ").title() for d in sorted_dims]
    scores = [d[1]["score"] for d in sorted_dims]
    colors = ["#27AE60" if s >= PASS_THRESHOLD else "#E74C3C" for s in scores]

    fig, ax = plt.subplots(figsize=(8, 6))
    bars = ax.barh(names, scores, color=colors, height=0.6, edgecolor="white", linewidth=0.5)
    ax.axvline(x=PASS_THRESHOLD, color="#2E86AB", linestyle="--", linewidth=1.5, label=f"Pass Threshold ({PASS_THRESHOLD})")
    ax.set_xlim(0, 1.05)
    ax.set_xlabel("Score", fontsize=11, fontweight="bold")
    ax.set_title("Dimension Scores Overview", fontsize=14, fontweight="bold", color="#1B2A4A", pad=15)
    ax.legend(loc="lower right", fontsize=9)
    ax.invert_yaxis()

    for bar, score in zip(bars, scores):
        ax.text(score + 0.02, bar.get_y() + bar.get_height()/2, f"{score:.2f}",
                va="center", fontsize=9, fontweight="bold")

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    plt.tight_layout()
    plt.savefig(filename, dpi=200, bbox_inches="tight")
    plt.close()


def _create_category_radar(dimensions: dict, filename: str):
    """Create a radar/spider chart by category."""
    categories: dict[str, list[float]] = {}
    for dim, info in dimensions.items():
        categories.setdefault(info["category"], []).append(info["score"])
    cat_means = {cat: sum(s)/len(s) for cat, s in categories.items()}

    labels = list(cat_means.keys())
    values = list(cat_means.values())
    N = len(labels)

    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    values += values[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, values, color="#2E86AB", alpha=0.25)
    ax.plot(angles, values, color="#2E86AB", linewidth=2)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=10, fontweight="bold")
    ax.set_ylim(0, 1.0)
    ax.set_title("Category Performance", fontsize=14, fontweight="bold", color="#1B2A4A", pad=20)

    # Add threshold circle
    threshold_values = [PASS_THRESHOLD] * (N + 1)
    ax.plot(angles, threshold_values, color="#E74C3C", linestyle="--", linewidth=1, alpha=0.7)

    plt.tight_layout()
    plt.savefig(filename, dpi=200, bbox_inches="tight")
    plt.close()


def _create_pass_fail_pie(dimensions: dict, filename: str):
    """Create a pass/fail pie chart."""
    passing = sum(1 for d in dimensions.values() if d["score"] >= PASS_THRESHOLD)
    failing = len(dimensions) - passing

    fig, ax = plt.subplots(figsize=(5, 5))
    wedges, texts, autotexts = ax.pie(
        [passing, failing],
        labels=["Pass", "Fail"],
        colors=["#27AE60", "#E74C3C"],
        autopct="%1.0f%%",
        startangle=90,
        textprops={"fontsize": 13, "fontweight": "bold"},
        wedgeprops={"edgecolor": "white", "linewidth": 2},
    )
    for t in autotexts:
        t.set_color("white")
        t.set_fontsize(14)
    ax.set_title(f"Pass/Fail Breakdown ({passing}/{len(dimensions)} dimensions)",
                 fontsize=13, fontweight="bold", color="#1B2A4A", pad=15)
    plt.tight_layout()
    plt.savefig(filename, dpi=200, bbox_inches="tight")
    plt.close()


def generate_report():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── Default font ──────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ══════════════════════════════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Agent Evaluation Report")
    run.font.size = Pt(32)
    run.font.color.rgb = BRAND_DARK
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Agent: {AGENT_NAME.replace('_', ' ').title()}")
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_MID
    run.bold = True

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Evaluation Date: {EVAL_DATE}\nFramework: wega-evals v0.1.0\nJudge Model: {JUDGE_MODEL}")
    run.font.size = Pt(11)
    run.font.color.rgb = GREY

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("─" * 60)
    run.font.color.rgb = BRAND_MID

    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = classification.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
    run.font.size = Pt(10)
    run.font.color.rgb = FAIL_RED
    run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  TABLE OF CONTENTS (placeholder)
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Executive Summary",
        "2. Evaluation Methodology",
        "3. Results Overview",
        "4. Detailed Dimension Analysis",
        "5. Category Performance",
        "6. Strengths & Weaknesses",
        "7. Recommendations",
        "8. Appendix",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "1. Executive Summary", level=1)

    all_scores = [d["score"] for d in DIMENSIONS.values()]
    overall = sum(all_scores) / len(all_scores)
    passing = sum(1 for s in all_scores if s >= PASS_THRESHOLD)
    failing = len(all_scores) - passing
    verdict = "PASS" if overall >= PASS_THRESHOLD else "NEEDS IMPROVEMENT"

    doc.add_paragraph(
        f"This report presents the evaluation results for the '{AGENT_NAME}' AI agent, "
        f"assessed across {len(DIMENSIONS)} quality dimensions using {ITEMS_EVALUATED} "
        f"production traces from the '{DATASET}' dataset."
    )

    doc.add_paragraph(AGENT_DESCRIPTION)

    # KPI Table
    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_table.style = "Table Grid"

    headers = ["Overall Score", "Verdict", "Dimensions Passing", "Items Evaluated"]
    values_row = [f"{overall:.2f}", verdict, f"{passing}/{len(DIMENSIONS)}", str(ITEMS_EVALUATED)]
    colors_row = [
        BRAND_MID.hex() if overall >= PASS_THRESHOLD else "E74C3C",
        "27AE60" if verdict == "PASS" else "E74C3C",
        "2E86AB",
        "2E86AB",
    ]

    for i, (h, v) in enumerate(zip(headers, values_row)):
        cell = kpi_table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            run.bold = True

        cell2 = kpi_table.cell(1, i)
        cell2.text = v
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.size = Pt(16)
            run.bold = True
            if i == 0:
                run.font.color.rgb = PASS_GREEN if overall >= PASS_THRESHOLD else FAIL_RED
            elif i == 1:
                run.font.color.rgb = PASS_GREEN if verdict == "PASS" else FAIL_RED
            else:
                run.font.color.rgb = BRAND_MID

    doc.add_paragraph()

    doc.add_paragraph(
        f"The agent achieved an overall score of {overall:.2f} against a pass threshold of "
        f"{PASS_THRESHOLD:.2f}. Out of {len(DIMENSIONS)} evaluated dimensions, "
        f"{passing} passed and {failing} failed. The primary areas of concern are "
        f"hallucination in generated scripts, low faithfulness to test case specifications, "
        f"slow response times, and overly verbose script output."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  2. EVALUATION METHODOLOGY
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "2. Evaluation Methodology", level=1)

    _add_styled_heading(doc, "2.1 Framework", level=2)
    doc.add_paragraph(
        "Evaluations were conducted using the wega-evals framework (v0.1.0), a reusable "
        "LLM agent evaluation system that combines LLM-as-judge assessments with "
        "programmatic metrics. All evaluation scores are stored in Langfuse for "
        "traceability and audit purposes."
    )

    _add_styled_heading(doc, "2.2 Evaluation Approaches", level=2)

    approach_table = doc.add_table(rows=3, cols=3)
    approach_table.style = "Table Grid"
    approach_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    approach_headers = ["Approach", "Dimensions", "Description"]
    for i, h in enumerate(approach_headers):
        cell = approach_table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

    llm_dims = [d for d, info in DIMENSIONS.items() if info["evaluator"] == "LLM Judge"]
    prog_dims = [d for d, info in DIMENSIONS.items() if info["evaluator"] == "Programmatic"]

    approaches = [
        ("LLM-as-Judge", str(len(llm_dims)),
         f"Gemini 2.5 Flash evaluates agent output against structured rubrics. "
         f"Each dimension has a custom prompt template producing a 0-1 score with reasoning."),
        ("Programmatic", str(len(prog_dims)),
         f"Rule-based metrics computed from trace metadata: latency (response time), "
         f"cost efficiency (token usage), and token efficiency."),
    ]
    for row_idx, (approach, count, desc) in enumerate(approaches, 1):
        approach_table.cell(row_idx, 0).text = approach
        approach_table.cell(row_idx, 1).text = count
        approach_table.cell(row_idx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        approach_table.cell(row_idx, 2).text = desc

    doc.add_paragraph()

    _add_styled_heading(doc, "2.3 Dataset", level=2)
    doc.add_paragraph(
        f"The evaluation dataset '{DATASET}' contains {ITEMS_EVALUATED} items seeded from "
        f"production Langfuse traces of type 'testcase_to_script_conversion'. Each item "
        f"includes the original test case input and the agent's generated test automation "
        f"script output."
    )

    _add_styled_heading(doc, "2.4 Scoring", level=2)
    doc.add_paragraph(
        f"All scores are normalized to a 0.0–1.0 scale. A dimension is considered passing "
        f"if its mean score across all items is ≥ {PASS_THRESHOLD:.2f}. The overall score "
        f"is the arithmetic mean of all dimension scores."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  3. RESULTS OVERVIEW
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "3. Results Overview", level=1)

    # Bar chart
    bar_chart_path = os.path.join(os.path.dirname(__file__), "_report_scripts_bar.png")
    _create_bar_chart(DIMENSIONS, bar_chart_path)
    doc.add_picture(bar_chart_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Full results table
    _add_styled_heading(doc, "3.1 Dimension Scores", level=2)

    sorted_dims = sorted(DIMENSIONS.items(), key=lambda x: x[1]["score"], reverse=True)
    results_table = doc.add_table(rows=len(sorted_dims) + 1, cols=5)
    results_table.style = "Table Grid"
    results_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_headers = ["Dimension", "Category", "Evaluator", "Score", "Status"]
    for i, h in enumerate(col_headers):
        cell = results_table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

    for row_idx, (dim, info) in enumerate(sorted_dims, 1):
        score = info["score"]
        status = "PASS" if score >= PASS_THRESHOLD else "FAIL"

        results_table.cell(row_idx, 0).text = dim.replace("_", " ").title()
        results_table.cell(row_idx, 1).text = info["category"]
        results_table.cell(row_idx, 2).text = info["evaluator"]

        score_cell = results_table.cell(row_idx, 3)
        score_cell.text = f"{score:.2f}"
        score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in score_cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = PASS_GREEN if score >= PASS_THRESHOLD else FAIL_RED

        status_cell = results_table.cell(row_idx, 4)
        status_cell.text = status
        status_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(status_cell, "27AE60" if status == "PASS" else "E74C3C")
        for run in status_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

        # Alternate row shading
        if row_idx % 2 == 0:
            for col in range(3):
                _set_cell_shading(results_table.cell(row_idx, col), "F5F7FA")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  4. DETAILED DIMENSION ANALYSIS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "4. Detailed Dimension Analysis", level=1)

    dim_descriptions = {
        "script_correctness": (
            "Measures whether the generated test scripts are syntactically correct and would "
            "execute without errors in the target framework (Playwright, Selenium BDD).",
            "At 0.57, roughly half of the generated scripts contain syntax errors, incorrect "
            "API usage, or broken imports. While the overall structure is usually sound, "
            "individual assertions and element selectors frequently have issues that would "
            "prevent successful execution without manual correction."
        ),
        "framework_compliance": (
            "Evaluates whether the generated scripts correctly follow the conventions and "
            "best practices of the specified test framework (e.g., Playwright test patterns, "
            "Selenium Page Object Model, BDD step definitions).",
            "Near-passing at 0.69, the agent demonstrates good framework awareness. Most "
            "scripts use appropriate framework constructs (describe/it blocks, page fixtures, "
            "Given/When/Then steps), though some outputs mix framework idioms or use deprecated APIs."
        ),
        "test_coverage_mapping": (
            "Checks whether all test steps and assertions from the source test case are "
            "faithfully represented in the generated script.",
            "At 0.48, the agent frequently misses test steps or omits critical assertions "
            "from the original test case. Negative scenarios and edge case validations are "
            "particularly under-represented in the generated scripts."
        ),
        "code_quality": (
            "Assesses the overall quality of the generated code including readability, "
            "maintainability, naming conventions, and adherence to coding standards.",
            "Scoring 0.66, the generated code is generally readable with reasonable variable "
            "naming. However, scripts often lack proper error handling, hardcode test data "
            "instead of parameterizing, and sometimes contain redundant or dead code blocks."
        ),
        "executability": (
            "Determines whether the generated scripts can be executed end-to-end in a "
            "properly configured test environment without modifications.",
            "At 0.54, roughly half of scripts would fail execution due to missing imports, "
            "undefined variables, incorrect locator strategies, or improper async/await usage. "
            "Scripts targeting Playwright tend to be more executable than Selenium ones."
        ),
        "coherence": (
            "Evaluates the logical flow and internal consistency of the generated test scripts.",
            "Passing at 0.73, scripts generally follow a logical setup-action-assertion flow. "
            "Test steps are ordered sensibly, and the relationship between actions and "
            "assertions is clear. This is the strongest output quality dimension."
        ),
        "conciseness": (
            "Measures whether scripts are appropriately sized without unnecessary verbosity, "
            "boilerplate, or over-engineering.",
            "Low at 0.37, the agent tends to produce overly verbose scripts with excessive "
            "comments, redundant helper functions, unnecessary abstraction layers, and "
            "boilerplate code that inflates script size without adding test value."
        ),
        "consistency": (
            "Checks whether the agent produces similar quality and style of scripts for "
            "similar test case inputs.",
            "At 0.62, there is noticeable variation in output quality across similar inputs. "
            "The same type of test case may produce well-structured scripts in one run and "
            "poorly organized ones in another, indicating non-deterministic generation patterns."
        ),
        "faithfulness": (
            "Evaluates whether the generated scripts accurately implement the test case "
            "specifications without deviating from the stated requirements.",
            "Critically low at 0.32, scripts frequently deviate from the original test case "
            "intent. The agent adds test scenarios not present in the spec, modifies assertion "
            "conditions, or changes the expected behavior being tested. This undermines the "
            "reliability of the generated scripts as a faithful translation of test cases."
        ),
        "hallucination": (
            "Detects fabricated code elements — non-existent APIs, invented selectors, "
            "made-up page URLs, or fictional framework methods not grounded in the input.",
            "Critical concern at 0.18. The agent frequently generates code referencing "
            "non-existent CSS selectors, invents page element IDs, fabricates API endpoints, "
            "and uses fictional framework methods. This is the most severe quality issue as "
            "hallucinated code elements make scripts fundamentally broken."
        ),
        "toxicity": (
            "Screens for harmful, offensive, or inappropriate content in generated scripts.",
            "Perfect score of 1.00. No toxic, offensive, or inappropriate content detected "
            "in any generated test scripts. Variable names, comments, and test descriptions "
            "are all professionally appropriate."
        ),
        "data_privacy_compliance": (
            "Checks that generated scripts don't contain or expose real PII, credentials, "
            "or sensitive information in test data, URLs, or comments.",
            "At 0.66, most scripts use synthetic test data. However, some outputs include "
            "patterns resembling real email addresses, phone numbers, or API keys in "
            "hardcoded test data, posing a potential data leakage risk if scripts are "
            "committed to repositories."
        ),
        "policy_compliance": (
            "Verifies adherence to organizational coding standards, test naming conventions, "
            "and automation framework usage policies.",
            "At 0.54, scripts partially follow organizational standards. Test naming conventions "
            "are inconsistently applied, some scripts use deprecated framework patterns, and "
            "file structure conventions are not always followed."
        ),
        "structural_validity": (
            "Verifies that the output follows the expected format — valid code blocks with "
            "proper language markers, correct file structure, and parseable syntax.",
            "Perfect score of 1.00. All outputs are well-structured with proper code block "
            "formatting, correct language identifiers, and valid overall document structure. "
            "This indicates the agent reliably produces properly formatted output containers."
        ),
        "latency": (
            "Measures response time efficiency (≤10s = 1.0, ≤30s = 0.8, ≤60s = 0.6, etc.).",
            "Low at 0.35. The agent frequently exceeds acceptable response time thresholds, "
            "with many traces taking over 30 seconds to generate scripts. Complex test cases "
            "with multiple steps are particularly slow, likely due to the verbose output style."
        ),
        "cost_efficiency": (
            "Evaluates token consumption efficiency relative to output quality "
            "(≤5K tokens = 1.0, ≤20K = 0.8).",
            "At 0.67, token usage is higher than ideal. The verbose generation style consumes "
            "more tokens than necessary, though it remains within a manageable range. "
            "Reducing output verbosity would directly improve this metric."
        ),
        "token_efficiency": (
            "Measures the ratio of useful output tokens to total tokens consumed, indicating "
            "how efficiently the model converts input context into valuable test script output.",
            "At 0.57, a significant portion of generated tokens are low-value boilerplate, "
            "excessive comments, or redundant code. Improving conciseness would directly "
            "boost token efficiency."
        ),
    }

    for dim, info in sorted_dims:
        score = info["score"]
        status = "PASS" if score >= PASS_THRESHOLD else "FAIL"
        desc, analysis = dim_descriptions.get(dim, ("No description available.", "No analysis available."))

        heading = doc.add_heading(level=3)
        run = heading.add_run(f"{dim.replace('_', ' ').title()}  ")
        run.font.color.rgb = BRAND_DARK

        badge_run = heading.add_run(f"  {score:.2f} — {status}  ")
        badge_run.font.size = Pt(11)
        badge_run.font.color.rgb = PASS_GREEN if status == "PASS" else FAIL_RED
        badge_run.bold = True

        doc.add_paragraph(desc).runs[0].font.italic = True
        doc.add_paragraph(analysis)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  5. CATEGORY PERFORMANCE
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "5. Category Performance", level=1)

    # Radar chart
    radar_path = os.path.join(os.path.dirname(__file__), "_report_scripts_radar.png")
    _create_category_radar(DIMENSIONS, radar_path)
    doc.add_picture(radar_path, width=Inches(4.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Category summary table
    categories: dict[str, list[float]] = {}
    for dim, info in DIMENSIONS.items():
        categories.setdefault(info["category"], []).append(info["score"])

    cat_table = doc.add_table(rows=len(categories) + 1, cols=4)
    cat_table.style = "Table Grid"
    cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(["Category", "Avg Score", "Dimensions", "Status"]):
        cell = cat_table.cell(0, i)
        cell.text = h
        _set_cell_shading(cell, "1B2A4A")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)

    for row_idx, (cat, scores) in enumerate(sorted(categories.items()), 1):
        mean = sum(scores) / len(scores)
        status = "PASS" if mean >= PASS_THRESHOLD else "FAIL"
        cat_table.cell(row_idx, 0).text = cat
        score_cell = cat_table.cell(row_idx, 1)
        score_cell.text = f"{mean:.2f}"
        score_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in score_cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = PASS_GREEN if mean >= PASS_THRESHOLD else FAIL_RED
        cat_table.cell(row_idx, 2).text = str(len(scores))
        cat_table.cell(row_idx, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        status_cell = cat_table.cell(row_idx, 3)
        status_cell.text = status
        status_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(status_cell, "27AE60" if status == "PASS" else "E74C3C")
        for run in status_cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  6. STRENGTHS & WEAKNESSES
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "6. Strengths & Weaknesses", level=1)

    # Pie chart
    pie_path = os.path.join(os.path.dirname(__file__), "_report_scripts_pie.png")
    _create_pass_fail_pie(DIMENSIONS, pie_path)
    doc.add_picture(pie_path, width=Inches(4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _add_styled_heading(doc, "6.1 Key Strengths", level=2)
    strengths = [
        ("Structural Validity (1.00)",
         "The agent consistently produces well-formatted output with proper code block structure, "
         "language markers, and parseable syntax. Output containers are always valid."),
        ("Toxicity (1.00)",
         "No harmful, offensive, or inappropriate content detected in any generated scripts. "
         "All variable names, comments, and descriptions are professionally appropriate."),
        ("Coherence (0.73)",
         "Generated scripts follow a logical setup-action-assertion flow. Test steps are ordered "
         "sensibly and the relationship between actions and assertions is clear and readable."),
        ("Framework Compliance (0.69)",
         "Near-passing score demonstrating good framework awareness. Most scripts correctly use "
         "Playwright or Selenium BDD patterns, fixtures, and assertion libraries."),
    ]
    for title, desc in strengths:
        p = doc.add_paragraph()
        run = p.add_run(f"✓ {title}: ")
        run.bold = True
        run.font.color.rgb = PASS_GREEN
        p.add_run(desc)

    _add_styled_heading(doc, "6.2 Critical Weaknesses", level=2)
    weaknesses = [
        ("Hallucination (0.18)",
         "The most critical issue — the agent frequently fabricates non-existent CSS selectors, "
         "invents element IDs, uses fictional framework methods, and references made-up API "
         "endpoints. This renders many scripts fundamentally broken and unexecutable."),
        ("Faithfulness (0.32)",
         "Scripts frequently deviate from the original test case specifications. The agent adds "
         "unrequested test scenarios, modifies assertion conditions, and changes expected behavior, "
         "undermining trust in the generated output."),
        ("Latency (0.35)",
         "Response times frequently exceed 30 seconds, particularly for complex multi-step test "
         "cases. The verbose generation style contributes to slow performance."),
        ("Conciseness (0.37)",
         "Scripts are overly verbose with excessive comments, redundant helper functions, "
         "unnecessary abstraction layers, and boilerplate that inflates size without test value."),
    ]
    for title, desc in weaknesses:
        p = doc.add_paragraph()
        run = p.add_run(f"✗ {title}: ")
        run.bold = True
        run.font.color.rgb = FAIL_RED
        p.add_run(desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  7. RECOMMENDATIONS
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "7. Recommendations", level=1)

    doc.add_paragraph(
        "Based on the evaluation results, the following improvements are recommended "
        "in order of priority:"
    )

    recommendations = [
        ("P0 — Address Hallucination in Script Generation",
         "The agent must be constrained to only reference elements, selectors, and APIs "
         "that are explicitly present in the input test case or a provided page object model. "
         "Implement grounding techniques such as retrieval-augmented generation (RAG) with "
         "actual application DOM snapshots, or require the agent to cite the source of every "
         "selector it uses. Consider adding a post-generation validation step that flags "
         "any selector or API call not traceable to the input context."),
        ("P0 — Improve Faithfulness to Test Case Specifications",
         "Add explicit chain-of-thought prompting that requires the agent to first enumerate "
         "all test steps from the input, then generate script code for each step with a "
         "traceability mapping. Implement a verification pass that cross-references generated "
         "assertions against the original test case expected results. Penalize deviation from "
         "the spec during fine-tuning or prompt optimization."),
        ("P1 — Optimize Latency and Token Usage",
         "Reduce output verbosity by instructing the agent to omit unnecessary comments, avoid "
         "redundant helper functions, and generate minimal-but-complete scripts. Consider "
         "streaming output for long scripts. Investigate whether a smaller, faster model can "
         "handle simpler test cases while routing complex ones to the full model."),
        ("P1 — Improve Conciseness",
         "Update the system prompt to explicitly discourage boilerplate code, excessive inline "
         "comments, and over-engineered abstractions. Provide few-shot examples of concise, "
         "focused test scripts that demonstrate the desired level of brevity. Set maximum "
         "output token limits per script complexity tier."),
        ("P2 — Improve Consistency Across Similar Inputs",
         "Standardize output templates for each framework (Playwright JS, Selenium Java BDD, "
         "etc.) to reduce variation across runs. Lower generation temperature for more "
         "deterministic output. Implement output normalization post-processing to enforce "
         "consistent naming conventions and code structure."),
        ("P2 — Enhance Test Coverage Mapping",
         "Implement a structured input format that enumerates test steps as a checklist, "
         "requiring the agent to produce a corresponding script section for each step. Add "
         "a coverage report in the output that maps each test case step to its generated "
         "script counterpart, making gaps immediately visible."),
    ]

    for title, desc in recommendations:
        _add_styled_heading(doc, title, level=3)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════
    #  8. APPENDIX
    # ══════════════════════════════════════════════════════════════════

    _add_styled_heading(doc, "8. Appendix", level=1)

    _add_styled_heading(doc, "A. Evaluation Configuration", level=2)
    config_table = doc.add_table(rows=8, cols=2)
    config_table.style = "Table Grid"
    config_items = [
        ("Framework", "wega-evals v0.1.0"),
        ("Judge Model", JUDGE_MODEL),
        ("Judge Temperature", "0.0"),
        ("Dataset", DATASET),
        ("Items Evaluated", str(ITEMS_EVALUATED)),
        ("Pass Threshold", str(PASS_THRESHOLD)),
        ("Evaluation Date", EVAL_DATE),
        ("Trace Source", "Langfuse (production)"),
    ]
    for i, (k, v) in enumerate(config_items):
        key_cell = config_table.cell(i, 0)
        key_cell.text = k
        for run in key_cell.paragraphs[0].runs:
            run.bold = True
        _set_cell_shading(key_cell, "F5F7FA")
        config_table.cell(i, 1).text = v

    doc.add_paragraph()

    _add_styled_heading(doc, "B. Scoring Scale", level=2)
    doc.add_paragraph(
        "All dimension scores use a 0.0–1.0 continuous scale:\n"
        "• 0.90–1.00: Excellent\n"
        "• 0.70–0.89: Good (Pass)\n"
        "• 0.50–0.69: Fair\n"
        "• 0.30–0.49: Poor\n"
        "• 0.00–0.29: Critical"
    )

    _add_styled_heading(doc, "C. Disclaimer", level=2)
    disclaimer = doc.add_paragraph(
        "This evaluation was performed using automated LLM-as-judge techniques and "
        "programmatic metrics. While the evaluation framework aims for objectivity, "
        "LLM-based assessments may contain inherent biases. Results should be interpreted "
        "in conjunction with human review and domain expertise. Scores reflect the quality "
        "of agent outputs at the time of evaluation and may vary with different input "
        "distributions or updated agent versions."
    )
    disclaimer.runs[0].font.italic = True
    disclaimer.runs[0].font.size = Pt(10)
    disclaimer.runs[0].font.color.rgb = GREY

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Generated by wega-evals • {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY

    # ── Save ──────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(__file__), f"{AGENT_NAME}_evaluation_report.docx")
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

    # Clean up temp chart images
    for path in [bar_chart_path, radar_path, pie_path]:
        if os.path.exists(path):
            os.remove(path)

    return output_path


if __name__ == "__main__":
    generate_report()
