"""
utils/file_extractor.py

Extracts plain text from uploaded .docx and .pdf files.
Used to turn user-uploaded attachments into text fed to the BRD generation LLM.
"""
from __future__ import annotations
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

# Magic-byte signatures for the supported binary formats.  TXT has no
# fixed signature, so we rely on a printable-character heuristic instead.
_MAGIC_PDF  = b"%PDF-"
_MAGIC_ZIP  = b"PK\x03\x04"      # docx is a zip container
_MAGIC_OLE  = b"\xD0\xCF\x11\xE0"  # legacy .doc (CFB/OLE2)


def verify_file_signature(filename: str, file_bytes: bytes) -> str | None:
    """Return ``None`` when the magic bytes match the filename extension.

    Otherwise return a human-readable error describing the mismatch.  This
    lets callers reject uploads where the extension lies about the content
    (e.g. a renamed ``.exe`` masquerading as ``.pdf``).
    """
    if not file_bytes:
        return "file is empty"
    ext = Path(filename).suffix.lower()
    head = file_bytes[:8]
    if ext == ".pdf":
        if not head.startswith(_MAGIC_PDF):
            return "content does not look like a PDF (missing %PDF- header)"
    elif ext == ".docx":
        if not head.startswith(_MAGIC_ZIP):
            return "content does not look like a .docx (missing ZIP header)"
    elif ext == ".doc":
        if not head.startswith(_MAGIC_OLE):
            return "content does not look like a legacy .doc file (missing OLE header)"
    elif ext == ".txt":
        # Reject obvious binary content masquerading as text.
        sample = file_bytes[:4096]
        if b"\x00" in sample:
            return "text file contains NUL bytes (likely binary)"
    return None


def extract_text_from_bytes(filename: str, file_bytes: bytes) -> str:
    """
    Extract plain text from a file given its original filename and raw bytes.
    Supports: .pdf, .docx, .doc, .txt

    Returns extracted text string, or raises ValueError for unsupported types
    or content-vs-extension mismatches.
    """
    ext = Path(filename).suffix.lower()

    sig_error = verify_file_signature(filename, file_bytes)
    if sig_error:
        raise ValueError(f"'{filename}': {sig_error}")

    if ext == ".pdf":
        return _extract_pdf(file_bytes, filename)
    elif ext == ".docx":
        return _extract_docx(file_bytes, filename)
    elif ext == ".doc":
        return _extract_doc(file_bytes, filename)
    elif ext == ".txt":
        return _extract_txt(file_bytes, filename)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Please upload .pdf, .docx, .doc, or .txt files only."
        )


def _clean_pdf_text(text: str) -> str:
    """
    Clean up common PDF extraction artefacts:
    - Replace mathematical italic Unicode letters with their ASCII equivalents
      so formulas like 𝑤ᵢ, 𝑣, 𝑓, 𝑝, 𝑀 come through as plain letters.
    - Strip stray PDF encoding junk that appears around inline formula fragments.
    """
    # Mathematical italic / bold-italic single letters (U+1D400 block)
    _MATH_ITALIC = str.maketrans(
        "𝑎𝑏𝑐𝑑𝑒𝑓𝑔ℎ𝑖𝑗𝑘𝑙𝑚𝑛𝑜𝑝𝑞𝑟𝑠𝑡𝑢𝑣𝑤𝑥𝑦𝑧"
        "𝐴𝐵𝐶𝐷𝐸𝐹𝐺𝐻𝐼𝐽𝐾𝐿𝑀𝑁𝑂𝑃𝑄𝑅𝑆𝑇𝑈𝑉𝑊𝑋𝑌𝑍",
        "abcdefghijklmnopqrstuvwxyz"
        "ABCDEFGHIJKLMNOPQRSTUVWXYZ",
    )
    text = text.translate(_MATH_ITALIC)
    # Unicode subscript digits/letters → plain
    _SUBS = str.maketrans("₀₁₂₃₄₅₆₇₈₉ᵢⱼₙ", "0123456789ijn")
    text = text.translate(_SUBS)
    # Unicode superscript digits → plain (used in exponents)
    _SUPS = str.maketrans("⁰¹²³⁴⁵⁶⁷⁸⁹ⁿ", "0123456789n")
    text = text.translate(_SUPS)
    # Unicode Sigma Σ is normally fine; leave it
    # Remove stray short junk tokens that appear around garbled inline formulas
    # (sequences like "!"# " or "$" that are encoding artefacts, not real content)
    import re
    text = re.sub(r'[!"#$%&\*]{1,4}\s+', "", text)
    return text


def _tables_to_markdown(page) -> list[str]:
    """
    Extract tables from a pdfplumber Page and return each as a markdown table string.
    Skips tables that are empty or have only one column (likely false positives).
    """
    md_tables = []
    try:
        tables = page.extract_tables()
    except Exception:
        return md_tables
    for table in tables:
        if not table:
            continue
        # Filter rows: skip entirely-None rows
        rows = [[cell or "" for cell in row] for row in table if any(cell for cell in row)]
        if len(rows) < 2:
            continue
        # Determine real column count (ignore trailing empty columns)
        max_cols = max(len(r) for r in rows)
        if max_cols < 2:
            continue
        # Pad rows
        rows = [r + [""] * (max_cols - len(r)) for r in rows]
        header = rows[0]
        separator = ["---"] * max_cols
        md_rows = [
            "| " + " | ".join(cell.replace("\n", " ").strip() for cell in row) + " |"
            for row in rows
        ]
        md_rows.insert(1, "| " + " | ".join(separator) + " |")
        md_tables.append("\n".join(md_rows))
    return md_tables


def _extract_pdf(data: bytes, filename: str) -> str:
    """Extract all text and tables from a PDF using pdfplumber.

    Tables found on each page are rendered as markdown pipe tables and
    prepended before the page's plain text so that calculation tables,
    requirement tables, and data schemas are clearly structured for the LLM.
    Mathematical Unicode artefacts from embedded formula fonts are cleaned up.
    """
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                page_sections: list[str] = []

                # 1. Tables — extracted and formatted as markdown first
                md_tables = _tables_to_markdown(page)
                if md_tables:
                    page_sections.append("**Tables on this page:**\n\n" + "\n\n".join(md_tables))

                # 2. Plain text — clean up math Unicode artefacts
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    page_sections.append(_clean_pdf_text(page_text.strip()))

                if page_sections:
                    text_parts.append(f"[Page {i}]\n" + "\n\n".join(page_sections))
        result = "\n\n".join(text_parts)
        if not result.strip():
            logger.warning("PDF '%s' yielded no extractable text (may be scanned/image-only).", filename)
            return f"[Note: '{filename}' appears to be a scanned PDF with no extractable text.]"
        logger.info("Extracted %d chars from PDF '%s'", len(result), filename)
        return result
    except Exception as exc:
        logger.error("PDF extraction failed for '%s': %s", filename, exc)
        raise ValueError(f"Could not extract text from PDF '{filename}': {exc}") from exc


def _extract_docx(data: bytes, filename: str) -> str:
    """Extract all text from a .docx file using python-docx.

    Includes paragraphs, tables, headers, and footers from every section.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts: list[str] = []

        # Headers / footers (per section) — captured up-front so the LLM
        # sees document-level boilerplate that often holds version, owner, etc.
        for sec_idx, section in enumerate(doc.sections, start=1):
            for label, source in (("header", section.header), ("footer", section.footer)):
                try:
                    text_chunks = [
                        p.text.strip() for p in source.paragraphs if p.text and p.text.strip()
                    ]
                except Exception:
                    text_chunks = []
                if text_chunks:
                    parts.append(f"[Section {sec_idx} {label}] " + " | ".join(text_chunks))

        # Body paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        result = "\n".join(parts)
        if not result.strip():
            return f"[Note: '{filename}' contained no extractable text.]"
        logger.info("Extracted %d chars from DOCX '%s'", len(result), filename)
        return result
    except Exception as exc:
        logger.error("DOCX extraction failed for '%s': %s", filename, exc)
        raise ValueError(f"Could not extract text from '{filename}': {exc}") from exc


def _extract_doc(data: bytes, filename: str) -> str:
    """Extract text from a legacy .doc file using pywin32 (Windows only).

    Uses Word COM automation to open the file and extract text content.
    Falls back to basic OLE text extraction if pywin32 is unavailable.
    """
    import tempfile
    import os
    import sys

    # Write bytes to a temp file (Word COM requires a file path)
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        tmp_path = tmp.name

    try:
        if sys.platform == "win32":
            try:
                import win32com.client
                import pythoncom

                pythoncom.CoInitialize()
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(tmp_path, ReadOnly=True)
                    try:
                        text = doc.Content.Text
                        result = text.strip() if text else ""
                        if not result:
                            return f"[Note: '{filename}' contained no extractable text.]"
                        logger.info("Extracted %d chars from DOC '%s' via Word COM", len(result), filename)
                        return result
                    finally:
                        doc.Close(False)
                        word.Quit()
                finally:
                    pythoncom.CoUninitialize()
            except ImportError:
                logger.warning("pywin32 not installed; falling back to basic OLE extraction for '%s'", filename)
                return _extract_doc_ole_fallback(data, filename)
            except Exception as exc:
                logger.warning("Word COM failed for '%s': %s; trying OLE fallback", filename, exc)
                return _extract_doc_ole_fallback(data, filename)
        else:
            # Non-Windows: use OLE fallback
            return _extract_doc_ole_fallback(data, filename)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _extract_doc_ole_fallback(data: bytes, filename: str) -> str:
    """Basic text extraction from .doc using olefile (OLE compound document).

    This is a fallback when pywin32/Word is unavailable. It extracts raw text
    from the WordDocument stream but may miss formatting and embedded objects.
    """
    try:
        import olefile
    except ImportError:
        raise ValueError(
            f"Cannot extract '{filename}': legacy .doc support requires "
            "either pywin32 (Windows with Word) or olefile package."
        )

    try:
        ole = olefile.OleFileIO(io.BytesIO(data))
        try:
            if not ole.exists("WordDocument"):
                raise ValueError(f"'{filename}' does not contain a valid Word document stream.")

            # Try to get text from the Word Document stream
            # The actual text is in a complex format; we'll try to extract readable portions
            word_stream = ole.openstream("WordDocument").read()

            # Extract printable ASCII/Unicode text chunks
            text_chunks = []
            current_chunk = []
            for byte in word_stream:
                if 32 <= byte < 127 or byte in (9, 10, 13):  # printable ASCII + whitespace
                    current_chunk.append(chr(byte))
                else:
                    if len(current_chunk) > 3:  # only keep chunks with 4+ chars
                        text_chunks.append("".join(current_chunk))
                    current_chunk = []
            if len(current_chunk) > 3:
                text_chunks.append("".join(current_chunk))

            result = " ".join(text_chunks)
            if not result.strip():
                return f"[Note: '{filename}' contained no extractable text (OLE fallback).]"
            logger.info("Extracted %d chars from DOC '%s' via OLE fallback", len(result), filename)
            return result
        finally:
            ole.close()
    except Exception as exc:
        logger.error("DOC extraction failed for '%s': %s", filename, exc)
        raise ValueError(f"Could not extract text from '{filename}': {exc}") from exc


def _extract_txt(data: bytes, filename: str) -> str:
    """Decode a plain text file, trying common encodings before falling back."""
    encodings = ("utf-8", "utf-8-sig", "utf-16", "cp1252", "latin-1", "gbk", "shift_jis")
    for encoding in encodings:
        try:
            text = data.decode(encoding)
            logger.info(
                "Extracted %d chars from TXT '%s' (encoding=%s)",
                len(text), filename, encoding,
            )
            return text
        except UnicodeDecodeError:
            continue
    # Last-resort: decode with replacement so we still surface *something*
    # rather than failing the whole upload.
    text = data.decode("utf-8", errors="replace")
    logger.warning(
        "TXT '%s' decoded with replacement characters — original encoding unknown.",
        filename,
    )
    return text


def format_extracted_docs(
    docs: list[tuple[str, str]]  # list of (filename, extracted_text)
) -> str:
    """
    Format multiple extracted documents into a single context block
    for the LLM prompt.
    """
    if not docs:
        return "No documents provided."
    parts = []
    for filename, text in docs:
        parts.append(
            f"=== Document: {filename} ===\n{text.strip()}\n=== End: {filename} ==="
        )
    return "\n\n".join(parts)