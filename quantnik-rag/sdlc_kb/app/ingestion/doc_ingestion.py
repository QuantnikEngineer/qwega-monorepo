import json
import re
from pathlib import Path
from typing import Any

import fitz  
import pandas as pd
from docx import Document as DocxDocument

from app.core.exceptions import UnsupportedFileTypeError, ParsingError


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".json"}


# ── Parsers ────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path) -> dict:
    try:
        doc = fitz.open(str(path))
        pages, sections = [], []
        for i, page in enumerate(doc):
            text = page.get_text("text")
            pages.append(text)
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") == 0:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            if span.get("size", 0) > 13:          # heading heuristic
                                title = span["text"].strip()
                                if title:
                                    sections.append({"title": title, "page": i + 1})
        return {
            "text":       "\n\n".join(pages),
            "sections":   sections,
            "page_count": len(pages),
            "metadata":   dict(doc.metadata),
        }
    except Exception as e:
        raise ParsingError(path.name, str(e))


def _parse_docx(path: Path) -> dict:
    try:
        doc = DocxDocument(str(path))
        paragraphs, sections, tables = [], [], []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
                if para.style.name.startswith("Heading"):
                    sections.append({"title": para.text.strip()})
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(rows)
        return {
            "text":       "\n".join(paragraphs),
            "sections":   sections,
            "page_count": 1,
            "metadata":   {"title": doc.core_properties.title or ""},
            "tables":     tables,
        }
    except Exception as e:
        raise ParsingError(path.name, str(e))


def _parse_xlsx(path: Path) -> dict:
    try:
        xls = pd.ExcelFile(str(path))
        parts, tables = [], []
        for sheet in xls.sheet_names:
            df = xls.parse(sheet).fillna("")
            parts.append(f"Sheet: {sheet}\n{df.to_string(index=False)}")
            tables.append({"sheet": sheet, "rows": df.to_dict(orient="records")})
        return {
            "text":       "\n\n".join(parts),
            "sections":   [{"title": s} for s in xls.sheet_names],
            "page_count": len(xls.sheet_names),
            "metadata":   {"sheets": xls.sheet_names},
            "tables":     tables,
        }
    except Exception as e:
        raise ParsingError(path.name, str(e))


def _parse_json(path: Path) -> dict:
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        def _flatten(obj, depth=0) -> str:
            pad = "  " * depth
            if isinstance(obj, dict):
                return "\n".join(f"{pad}{k}: {_flatten(v, depth+1)}" for k, v in obj.items())
            if isinstance(obj, list):
                return "\n".join(_flatten(i, depth) for i in obj)
            return f"{pad}{obj}"

        return {
            "text":       _flatten(data),
            "sections":   [],
            "page_count": 1,
            "metadata":   {"keys": list(data.keys()) if isinstance(data, dict) else []},
            "tables":     [],
        }
    except Exception as e:
        raise ParsingError(path.name, str(e))


# ── Router ─────────────────────────────────────────────────────────────────────

_PARSERS = {
    ".pdf":  _parse_pdf,
    ".docx": _parse_docx,
    ".xlsx": _parse_xlsx,
    ".json": _parse_json,
}


def parse(path: Path) -> dict:
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(ext)
    return _PARSERS[ext](path)


# ── Normalizer ─────────────────────────────────────────────────────────────────

def normalize(raw: dict) -> dict:
    text = raw.get("text", "")
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    text = text.strip()

    return {
        "text":       text,
        "sections":   raw.get("sections", []),
        "tables":     raw.get("tables", []),
        "page_count": raw.get("page_count", 1),
        "word_count": len(text.split()),
        "metadata":   raw.get("metadata", {}),
    }


# ── Public API ─────────────────────────────────────────────────────────────────

def ingest(path: Path) -> dict:
    """Single entry point: parse + normalize. Returns unified document dict."""
    raw = parse(path)
    return normalize(raw)